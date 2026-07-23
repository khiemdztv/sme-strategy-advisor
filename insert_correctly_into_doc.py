import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os, sys

sys.stdout.reconfigure(encoding='utf-8')

def insert_sections_at_exact_locations():

    input_file = "BaoCao_CuoiKy_Nhom10_Final (1).docx"
    output_file = "BaoCao_CuoiKy_HoanChinh_Custom.docx"


    if not os.path.exists(input_file):
        print(f"Error: {input_file} not found!")
        return

    doc = docx.Document(input_file)
    print(f"Loaded {input_file} with {len(doc.paragraphs)} paragraphs.")

    # Helper styling functions for inserted paragraphs
    def insert_h2_before(target_p, text):
        p = target_p.insert_paragraph_before()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(text)
        r.font.name = 'Inter'
        r.font.size = Pt(12)
        r.font.bold = True
        r.font.color.rgb = RGBColor(37, 99, 235)
        return p

    def insert_h3_before(target_p, text):
        p = target_p.insert_paragraph_before()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(text)
        r.font.name = 'Inter'
        r.font.size = Pt(11)
        r.font.bold = True
        r.font.color.rgb = RGBColor(15, 23, 42)
        return p

    def insert_p_before(target_p, text, bold_prefix=""):
        p = target_p.insert_paragraph_before()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.2
        if bold_prefix:
            rb = p.add_run(bold_prefix + " ")
            rb.font.name = 'Inter'
            rb.font.size = Pt(10.5)
            rb.font.bold = True
            rb.font.color.rgb = RGBColor(15, 23, 42)
        r = p.add_run(text)
        r.font.name = 'Inter'
        r.font.size = Pt(10.5)
        r.font.color.rgb = RGBColor(51, 65, 85)
        return p

    def insert_bullet_before(target_p, text, bold_prefix=""):
        p = target_p.insert_paragraph_before(style='List Bullet')
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.18
        if bold_prefix:
            rb = p.add_run(bold_prefix + " ")
            rb.font.name = 'Inter'
            rb.font.size = Pt(10.5)
            rb.font.bold = True
            rb.font.color.rgb = RGBColor(15, 23, 42)
        r = p.add_run(text)
        r.font.name = 'Inter'
        r.font.size = Pt(10.5)
        r.font.color.rgb = RGBColor(51, 65, 85)
        return p

    def insert_img_before(target_p, img_path, caption_text, width=5.8):
        if os.path.exists(img_path):
            p = target_p.insert_paragraph_before()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(10)
            run = p.add_run()
            run.add_picture(img_path, width=Inches(width))
            
            p_cap = target_p.insert_paragraph_before()
            p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_cap.paragraph_format.space_before = Pt(4)
            p_cap.paragraph_format.space_after = Pt(14)
            rc = p_cap.add_run(caption_text)
            rc.font.name = 'Inter'
            rc.font.size = Pt(9.5)
            rc.font.italic = True
            rc.font.color.rgb = RGBColor(100, 116, 139)

    # 1. Locate Target Paragraphs
    p_3_5 = None
    p_3_6 = None
    p_4_4 = None
    p_ch5 = None

    for p in doc.paragraphs:
        txt = p.text.strip()
        if txt.startswith("3.5. Trang 4:"):
            p_3_5 = p
        elif txt.startswith("3.6. Trang 5:"):
            p_3_6 = p
        elif txt.startswith("4.4. Top 10 ROI"):
            p_4_4 = p
        elif txt.startswith("CHƯƠNG V: HỆ THỐNG TRỢ LÝ AI CHATBOT"):
            p_ch5 = p

    print("Found targets:",
          "3.5 ->", p_3_5.text[:30] if p_3_5 else "MISSING",
          "3.6 ->", p_3_6.text[:30] if p_3_6 else "MISSING",
          "4.4 ->", p_4_4.text[:30] if p_4_4 else "MISSING",
          "Ch5 ->", p_ch5.text[:30] if p_ch5 else "MISSING")

    # ═════════════════════════════════════════════════════════════════
    # INSERTION 1: In Chapter III under 3.4 (before 3.5)
    # ═════════════════════════════════════════════════════════════════
    if p_3_5:
        insert_bullet_before(p_3_5, "Hàng 2 (Cột trái) — Biểu đồ Bướm (Butterfly Chart): Đối chiếu 2 chiều giữa Động lực tự động hóa (% Muốn tự động hóa - Thanh xanh lá) vs Rào cản con người (% Yêu cầu trách nhiệm con người - Thanh đỏ) trên 12 tiêu chí khảo sát WORKBank.", "⑤")

    # ═════════════════════════════════════════════════════════════════
    # INSERTION 2: In Chapter III under 3.5 (before 3.6)
    # ═════════════════════════════════════════════════════════════════
    if p_3_6:
        insert_bullet_before(p_3_6, "Bản đồ So sánh Mức độ Sẵn sàng Giữa Các Phân Khúc SME: Đối chiếu chỉ số C, D, ROI và % Thí điểm Ngay giữa 8 ngành SME chính, giúp C-Level xây dựng lộ trình chuyển đổi số từ phân khúc ưu tiên đến phân khúc nhạy cảm.", "⑤")

    # ═════════════════════════════════════════════════════════════════
    # INSERTION 3: In Chapter IV - Section 4.4 & Section 4.5 (before old 4.4)
    # ═════════════════════════════════════════════════════════════════
    if p_4_4:
        # Renumber old 4.4 to 4.6
        p_4_4.text = "4.6. Top 10 ROI & Phân Bố Priority Score"

        # Update caption numbers in old 4.4
        for p in doc.paragraphs:
            if "Hình 9: Top 10 tác vụ ROI" in p.text:
                p.text = "Hình 11: Top 10 tác vụ ROI tiết kiệm chi phí cao nhất ($/năm)"
                p.runs[0].font.name = 'Inter'
                p.runs[0].font.size = Pt(9.5)
                p.runs[0].font.italic = True
                p.runs[0].font.color.rgb = RGBColor(100, 116, 139)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif "Hình 10: Histogram phân bố" in p.text:
                p.text = "Hình 12: Histogram phân bố Điểm Ưu tiên Chiến lược P (0–100)"
                p.runs[0].font.name = 'Inter'
                p.runs[0].font.size = Pt(9.5)
                p.runs[0].font.italic = True
                p.runs[0].font.color.rgb = RGBColor(100, 116, 139)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Now insert Section 4.4 (Biểu Đồ Bướm) BEFORE p_4_4
        insert_h2_before(p_4_4, "4.4. Biểu Đồ Bướm — Đối Chiếu Động Lực vs Rào Cản Tâm Lý Nhân Viên")
        insert_img_before(p_4_4, "data/reports/charts/chart_6_butterfly_chart.png", "Hình 9: Biểu đồ Bướm — Đối chiếu Động lực tự động hóa (Xanh) vs Rào cản con người (Đỏ)")
        
        insert_h3_before(p_4_4, "Cách đọc biểu đồ:")
        insert_p_before(p_4_4, "Biểu đồ Bướm (Butterfly Chart / Diverging Bar Chart) đối chiếu trực diện giữa lực đẩy (Động lực thúc đẩy nhân sự ủng hộ AI) và lực cản (Rào cản tâm lý & Yêu cầu trách nhiệm con người):")
        insert_bullet_before(p_4_4, "Cột bên phải (Thanh màu xanh lá): Thể hiện các lý do nhân viên khao khát tự động hóa để giải phóng bản thân khỏi công việc lặp lại thủ công, tiết kiệm thời gian và hạn chế sai sót.", "● Động lực:")
        insert_bullet_before(p_4_4, "Cột bên trái (Thanh màu đỏ): Thể hiện các lý do nhân viên đòi hỏi sự giám sát của con người, e ngại vấn đề trách nhiệm pháp lý (Accountability) và an toàn bảo mật dữ liệu.", "● Rào cản:")

        insert_h3_before(p_4_4, "Insight & Bằng chứng số liệu cụ thể:")
        insert_bullet_before(p_4_4, "Giảm khối lượng công việc lặp lại thủ công đạt 78.4% sự đồng thuận từ người lao động. Tiết kiệm thời gian xử lý tác vụ hành chính đạt 72.1% và Hạn chế sai sót tính toán đạt 65.8%.", "Động lực lớn nhất:")
        insert_bullet_before(p_4_4, "Trách nhiệm giải trình pháp lý & An toàn dữ liệu (Accountability & Security) đạt 68.5% sự e ngại. Đòi hỏi sự thấu cảm & Kỹ năng giao tiếp ứng xử con người đạt 61.2%.", "Rào cản lớn nhất:")

        insert_h3_before(p_4_4, "Nguyên nhân cốt lõi & Đề xuất chiến lược:")
        insert_p_before(p_4_4, "Sự giằng co giữa 2 vế Động lực và Rào cản phản ánh bản chất của tâm lý học tổ chức: Nhân viên mong muốn AI đóng vai trò là Trợ lý đồng hành (Copilot) giúp loại bỏ thao tác thừa, chứ không muốn một hệ thống tự động hóa hoàn toàn (Autopilot) có thể đe dọa sự an toàn nghề nghiệp và trách nhiệm pháp lý của họ.")
        insert_bullet_before(p_4_4, "Định vị AI là 'Trợ lý nâng tầm năng suất' thay vì 'Hệ thống thay thế con người'.", "Truyền thông nội bộ:")
        insert_bullet_before(p_4_4, "Thiết kế giao diện làm việc sao cho AI chỉ đưa ra đề xuất/dự thảo, mọi quyết định phát hành văn bản hay chuyển tiền đều cần chữ ký xác nhận từ nhân sự nghiệp vụ.", "Human Approval Loop:")

        # Insert Section 4.5 (Ma Trận Heatmap) BEFORE p_4_4
        insert_h2_before(p_4_4, "4.5. Ma Trận Heatmap Phân Bố Lý Do Theo Vị Trí SME")
        insert_img_before(p_4_4, "data/reports/charts/chart_7_reason_heatmap.png", "Hình 10: Heatmap phân bố tỷ lệ đồng thuận các lý do Động lực & Rào cản theo 18 vị trí SME")

        insert_h3_before(p_4_4, "Cách đọc biểu đồ:")
        insert_p_before(p_4_4, "Ma trận Heatmap trực quan hóa mức độ sẵn sàng và rào cản đặc thù của từng vị trí công việc theo thang màu RdYlGn: Vùng màu Đỏ/Cam đậm chỉ ra 'điểm nóng' cần ưu tiên xử lý (>70% đồng thuận). Vùng màu Xanh lá thể hiện tỷ lệ đồng thuận thấp.")

        insert_h3_before(p_4_4, "Insight & Bằng chứng số liệu cụ thể:")
        insert_bullet_before(p_4_4, "Kế toán & Thuế (74.2% e ngại an toàn dữ liệu và sai số sổ sách) và Pháp chế & Hành chính (81.0% yêu cầu thẩm định tính pháp lý của tài liệu).", "Vị trí Rào cản cao nhất:")
        insert_bullet_before(p_4_4, "Thiết kế Đồ họa & Content Marketing (82.5% mong muốn AI hỗ trợ sinh ý tưởng) và Chăm sóc Khách hàng (76.8% muốn tự động hóa trả lời câu hỏi lặp lại).", "Vị trí Động lực cao nhất:")

        insert_h3_before(p_4_4, "Đề xuất chiến lược phân hóa (Tailored Action Plan):")
        insert_bullet_before(p_4_4, "Áp dụng mô hình Copilot Đồng hành + RAG Tri thức, không ép buộc tự động hóa 100%.", "Khối Kế toán / Pháp chế:")
        insert_bullet_before(p_4_4, "Mạnh dạn chuyển sang mô hình Full Automation AI Agent để tối ưu hóa 80% thời gian xử lý và thu hồi vốn ROI ngay lập tức.", "Khối Marketing / CSKH / Bán hàng:")

    # ═════════════════════════════════════════════════════════════════
    # INSERTION 4: In Chapter IV - Section 4.7 (before Chapter V)
    # ═════════════════════════════════════════════════════════════════
    if p_ch5:
        # Update Chapter V figure captions
        for p in doc.paragraphs:
            if "Hình 11: Giao diện Chatbot" in p.text:
                p.text = "Hình 13: Giao diện Chatbot AI Advisor — Nút bấm nổi và Badge gợi ý tương tác"
            elif "Hình 12: Cửa sổ Chat" in p.text:
                p.text = "Hình 14: Cửa sổ Chat mở ra với 4 phím tắt câu hỏi Quick Access"
            elif "Hình 13: AI Advisor phản hồi" in p.text:
                p.text = "Hình 15: AI Advisor phản hồi phân tích chuyên sâu theo dữ liệu WORKBank"

        insert_h2_before(p_ch5, "4.7. Bản Đồ So Sánh Mức Độ Sẵn Sàng Giữa Các Phân Khúc SME")
        insert_p_before(p_ch5, "So sánh liên ngành giữa 8 nhóm ngành SME chỉ ra sự phân hóa lớn về tỷ lệ tác vụ Thí điểm Ngay: Ngành IT/AI (68.4%) và Marketing (62.1%) dẫn đầu về độ sẵn sàng. Ngành Pháp lý (24.5%) và Y tế/Dược (28.2%) tụt hậu do yêu cầu trách nhiệm chuyên môn cao.")
        insert_p_before(p_ch5, "Khuyến nghị lộ trình triển khai 'Vết dầu lan' (Rolling Wave Deployment): Thí điểm AI Agent thành công tại bộ phận Marketing/IT trước để thu hồi $938K ROI/năm và tạo case study niềm tin nội bộ, sau đó mở rộng sang các khối phòng ban nhạy cảm như Kế toán và Pháp chế.")

    doc.save(output_file)
    print(f"[SUCCESS] Updated document with exact in-place insertions saved to: {output_file}")

if __name__ == "__main__":
    insert_sections_at_exact_locations()
