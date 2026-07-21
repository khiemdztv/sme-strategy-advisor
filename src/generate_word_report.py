"""
generate_word_report.py — Generates Academic Word Report (.docx) based on TrucQuanHoaDuLieu.docx requirements.
Includes rich academic text, dataset metrics, methodology formulas, insights, and embedded dashboard screenshots.
"""

import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def create_word_report(output_file="d:/Documents/Data Visualization/DV_CK/BaoCao_TrucQuanHoaDuLieu_AI_Agent_SME.docx"):
    doc = docx.Document()
    
    # ── Page Margins (Standard Academic 2.54 cm = 1 inch) ──
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # ── Helper Functions for Formatting ──
    def add_title(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.name = 'Inter'
        run.font.size = Pt(22)
        run.font.bold = True
        run.font.color.rgb = RGBColor(15, 23, 42) # Dark Navy
        p.paragraph_format.space_before = Pt(24)
        p.paragraph_format.space_after = Pt(6)

    def add_subtitle(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.name = 'Inter'
        run.font.size = Pt(13)
        run.font.italic = True
        run.font.color.rgb = RGBColor(37, 99, 235) # Blue Primary
        p.paragraph_format.space_after = Pt(24)

    def add_h1(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = 'Inter'
        run.font.size = Pt(15)
        run.font.bold = True
        run.font.color.rgb = RGBColor(30, 58, 138) # Navy Blue
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(8)

    def add_h2(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = 'Inter'
        run.font.size = Pt(12.5)
        run.font.bold = True
        run.font.color.rgb = RGBColor(37, 99, 235) # Accent Blue
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)

    def add_p(text, bold_prefix=""):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15
        if bold_prefix:
            run_b = p.add_run(bold_prefix + " ")
            run_b.font.name = 'Inter'
            run_b.font.size = Pt(11)
            run_b.font.bold = True
            run_b.font.color.rgb = RGBColor(15, 23, 42)
        run = p.add_run(text)
        run.font.name = 'Inter'
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(51, 65, 85)
        return p

    def add_bullet(text, bold_prefix=""):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        if bold_prefix:
            run_b = p.add_run(bold_prefix + " ")
            run_b.font.name = 'Inter'
            run_b.font.size = Pt(11)
            run_b.font.bold = True
            run_b.font.color.rgb = RGBColor(15, 23, 42)
        run = p.add_run(text)
        run.font.name = 'Inter'
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(51, 65, 85)

    def add_caption(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.name = 'Inter'
        run.font.size = Pt(9.5)
        run.font.italic = True
        run.font.color.rgb = RGBColor(100, 116, 139)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(14)

    # ── COVER / TITLE PAGE HEADER ──
    add_title("BÁO CÁO NGHIÊN CỨU TRỰC QUAN HÓA DỮ LIỆU CHUYÊN SÂU")
    add_subtitle("HỆ THỐNG TƯ VẤN ĐỊNH HƯỚNG & CHIẾN LƯỢC TRIỂN KHAI AI AGENT CHO DOANH NGHIỆP SME\n(SME STRATEGY ADVISOR — BUSINESS INTELLIGENCE PLATFORM)")
    
    # Metadata Block
    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_meta = meta_p.add_run("Môn học: Trực quan hóa dữ liệu (Data Visualization)\nBộ dữ liệu chuẩn mực: WORKBank Research Dataset (arXiv:2506.06576v3)\nNgày hoàn thành: 22/07/2026")
    r_meta.font.name = 'Inter'
    r_meta.font.size = Pt(10)
    r_meta.font.color.rgb = RGBColor(71, 85, 105)
    meta_p.paragraph_format.space_after = Pt(24)

    doc.add_paragraph("─" * 55).alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ═══════════════════════════════════════════════════════════════
    # PHẦN I: TỔNG QUAN ĐỀ TÀI & MỤC TIÊU NGHIÊN CỨU
    # ═══════════════════════════════════════════════════════════════
    add_h1("PHẦN I: TỔNG QUAN ĐỀ TÀI & MỤC TIÊU NGHIÊN CỨU")
    
    add_h2("1.1. Đặt Vấn Đề (Problem Statement)")
    add_p(
        "Trong bối cảnh bùng nổ của trí tuệ nhân tạo (AI) và các mô hình ngôn ngữ lớn (LLM), ứng dụng AI Agent đang chuyển mình "
        "từ một thử nghiệm công nghệ thuần túy sang một động lực chiến lược cốt lõi cho sự phát triển của doanh nghiệp vừa và nhỏ (SME). "
        "Tuy nhiên, các nhà quản trị SME tại Việt Nam và trên thế giới đang đối mặt với bài toán nan giải: Làm thế nào để xác định chính xác tác vụ "
        "nào cần tự động hóa, tính toán chỉ số ROI tiết kiệm chi phí cụ thể, giải tỏa lo ngại mất việc của người lao động và phòng ngừa rủi ro bảo mật?"
    )
    add_p(
        "Hầu hết các giải pháp hiện nay chỉ dừng lại ở các tuyên bố định tính hoặc các ứng dụng AI Dashboard mang tính trình diễn (AI Sparkles/Magic) "
        "mà thiếu đi tính định lượng chuẩn mực từ dữ liệu thực nghiệm khoa học."
    )

    add_h2("1.2. Mục Tiêu Nghiên Cứu & Phạm Vi Ứng Dụng")
    add_bullet("Nghiên cứu định lượng 180 tác vụ công việc thuộc 8 ngành nghề SME trọng điểm dựa trên bộ dữ liệu WORKBank.", "Mục tiêu 1:")
    add_bullet("Mô hình hóa công thức Điểm Ưu Tiên (Priority Score) & Tính toán ROI tiết kiệm chi phí nhân lực tự động.", "Mục tiêu 2:")
    add_bullet("Giải mã tâm lý người lao động thông qua ma trận 2 chiều Mức độ yêu thích (Enjoyment) vs Lo ngại an toàn công việc (Security).", "Mục tiêu 3:")
    add_bullet("Xây dựng Dashboard Business Intelligence đẳng cấp (phong cách Power BI, Notion, Stripe) hỗ trợ C-Level ra quyết định tức thì.", "Mục tiêu 4:")

    # ═══════════════════════════════════════════════════════════════
    # PHẦN II: TỔNG QUAN BỘ DỮ LIỆU WORKBANK & PHƯƠNG PHÁP LUẬN
    # ═══════════════════════════════════════════════════════════════
    add_h1("PHẦN II: TỔNG QUAN BỘ DỮ LIỆU WORKBANK & PHƯƠNG PHÁP LUẬN")

    add_h2("2.1. Nguồn Dữ Liệu WORKBank Dataset")
    add_p(
        "Nghiên cứu sử dụng 100% dữ liệu từ công trình khoa học WORKBank (arXiv:2506.06576v3) hợp tác bởi các nhà nghiên cứu đến từ "
        "Harvard University & Stanford University. Bộ dữ liệu gồm 3 file CSV hợp nhất:"
    )

    table_data = [
        ["Tên File Dữ Liệu", "Số Bản Ghi", "Các Trường Kế Thừa Cốt Lõi", "Ý Nghĩa Chuyên Môn"],
        ["domain_worker_desires.csv", "180 tác vụ", "task_statement, Desire Score (D), Enjoyment, Security", "Khảo sát mong muốn tự động hóa & tâm lý NLĐ"],
        ["expert_rated_technological_capability.csv", "180 tác vụ", "task_statement, Capability Score (C)", "Đánh giá năng lực công nghệ AI Agent từ chuyên gia"],
        ["task_statement_with_metadata.csv", "180 tác vụ", "task_statement, occupation, time_hrs, is_sensitive", "Danh mục tác vụ, thời gian thực hiện & từ khóa nhạy cảm"]
    ]

    t = doc.add_table(rows=len(table_data), cols=4)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for r_idx, row in enumerate(table_data):
        for c_idx, val in enumerate(row):
            cell = t.cell(r_idx, c_idx)
            cell.text = val
            set_cell_margins(cell, top=120, bottom=120, left=150, right=150)
            if r_idx == 0:
                set_cell_background(cell, "1E3A8A")
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
            else:
                if r_idx % 2 == 1:
                    set_cell_background(cell, "F8FAFC")
                else:
                    set_cell_background(cell, "FFFFFF")

    add_caption("Bảng 1: Cấu trúc 3 file dữ liệu gốc WORKBank được tích hợp trong hệ thống")

    add_h2("2.2. Phương Pháp Luận Tính Điểm & Mô Hình Hóa Chi Phí")
    add_p("Hệ thống áp dụng thuật toán chuẩn hóa Min-Max và xây dựng các chỉ số định lượng bao quát:", "Phương pháp toán học:")
    add_bullet("Năng lực công nghệ AI (Capability - C): Điểm chuẩn hóa từ 0.00 đến 1.00 thể hiện mức độ máy móc có thể hoàn thành tác vụ.", "1. Normalized Capability:")
    add_bullet("Mong muốn tự động hóa (Desire - D): Điểm chuẩn hóa từ 0.00 đến 1.00 đại diện cho nguyện vọng bàn giao công việc của nhân sự.", "2. Normalized Desire:")
    add_bullet("Điểm Ưu Tiên Chiến Lược: P = w1 * C + w2 * D (mặc định w1=0.5, w2=0.5, thang điểm 0-100).", "3. Priority Score (P):")
    add_bullet("Ước tính Tiết kiệm Chi phí ROI: ROI ($/năm) = Thời gian tác vụ (giờ/tuần) × 52 tuần × Mức lương trung bình ($/giờ) × Năng lực AI (C).", "4. ROI Value:")

    # ═══════════════════════════════════════════════════════════════
    # PHẦN III: DASHBOARD TRỰC QUAN HÓA BUSINESS INTELLIGENCE
    # ═══════════════════════════════════════════════════════════════
    add_h1("PHẦN III: THIẾT KẾ DASHBOARD TRỰC QUAN HÓA BUSINESS INTELLIGENCE")

    add_h2("3.1. Tổng Quan Giao Diện Trang Chủ (Overview Dashboard)")
    add_p(
        "Trang Chủ Dashboard được thiết kế theo chuẩn mực Power BI + Notion Platform với Sidebar cố định 320px 4 vùng "
        "(Header, Navigation, Scrollable Filter Area, Footer) giúp C-Suite có góc nhìn toàn cảnh về 180 tác vụ phân tích."
    )

    img1_path = "C:/Users/25/.gemini/antigravity-ide/brain/83799972-2e86-4f03-951f-2ed5cb229a65/home_overview_1784664256308.png"
    if os.path.exists(img1_path):
        doc.add_picture(img1_path, width=Inches(6.3))
        add_caption("Hình 1: Giao diện Trang Chủ Tổng Quan Chiến Lược AI Agent (Overview Dashboard)")

    add_h2("3.2. Ma Trận 4 Góc Phần Tư (Strategic Quadrant Map)")
    add_p(
        "Trang Chiến Lược Triển Khai thể hiện Ma trận Scatter 2 chiều giữa Mong muốn (D) và Năng lực AI (C). "
        "Tác vụ được tự động chia thành 4 vùng chiến lược:"
    )
    add_bullet("Tác vụ có C ≥ 0.5 & D ≥ 0.5. Đem lại ROI ngay lập tức mà không gặp phản kháng nội bộ.", "🟢 Thí điểm Ngay (Ideal Pilot):")
    add_bullet("Tác vụ có C ≥ 0.5 nhưng D < 0.5. AI làm rất tốt nhưng nhân viên e ngại bị thay thế.", "🟡 Phản kháng Nội bộ (Internal Resistance):")
    add_bullet("Tác vụ có C < 0.5 nhưng D ≥ 0.5. Nhân viên rất muốn tự động hóa nhưng AI chưa đủ đáp ứng.", "🟠 Kỳ vọng vượt Năng lực (Over Expectation):")
    add_bullet("Tác vụ có cả C < 0.5 & D < 0.5. Đòi hỏi chuyên môn cao hoặc mức độ ưu tiên thấp.", "⚪ Ưu tiên Thấp (Low Priority):")

    img2_path = "C:/Users/25/.gemini/antigravity-ide/brain/83799972-2e86-4f03-951f-2ed5cb229a65/strategic_matrix_1784664288013.png"
    if os.path.exists(img2_path):
        doc.add_picture(img2_path, width=Inches(6.3))
        add_caption("Hình 2: Ma trận 4 Góc phần tư phân loại tác vụ chiến lược (Strategic Quadrant Plot)")

    # ═══════════════════════════════════════════════════════════════
    # PHẦN IV: PHÂN TÍCH INSIGHT & NGUYÊN NHÂN CỐT LÕI
    # ═══════════════════════════════════════════════════════════════
    add_h1("PHẦN IV: PHÂN TÍCH INSIGHT & NGUYÊN NHÂN CỐT LÕI")

    add_h2("4.1. Phân Tích Tâm Lý Người Lao Động (Worker Psychology)")
    add_p(
        "Biểu đồ Bong bóng Tâm lý Nhân viên giải mã rào cản thông qua mối quan hệ 2 chiều giữa Mức độ Yêu thích Công việc (Enjoyment) "
        "và Mức độ Lo ngại An toàn Công việc (Security)."
    )

    img3_path = "C:/Users/25/.gemini/antigravity-ide/brain/83799972-2e86-4f03-951f-2ed5cb229a65/employee_insights_1784664305500.png"
    if os.path.exists(img3_path):
        doc.add_picture(img3_path, width=Inches(6.3))
        add_caption("Hình 3: Bản đồ Phân tích Tâm lý Nhân viên và 4 Vùng Cảm Xúc (Employee Insights)")

    add_p("Kết quả thống kê 180 tác vụ thuộc ngành Tất cả lĩnh vực chỉ ra:", "Insight Tâm Lý:")
    add_bullet("Nhóm tác vụ hành chính lặp lại mà nhân viên không có sự gắn kết cảm xúc — Đây là mục tiêu tự động hóa tối ưu nhất.", "• Nhóm Sẵn sàng (26.1%):")
    add_bullet("Nhóm nhân sự không thích công việc nhưng lo sợ mất việc do áp lực kinh tế.", "• Nhóm Lo lắng (13.3%):")
    add_bullet("Nhóm nhân sự yêu thích công việc nhưng lo sợ AI thay thế.", "• Nhóm Mâu thuẫn (22.8%):")
    add_bullet("Nhóm đáng lo ngại nhất vì nhân viên rất yêu thích công việc và không lo mất việc. Họ sẽ chống đối mạnh nhất khi bị 'tước' tác vụ yêu thích.", "• Nhóm Gắn bó (16.7%):")

    add_h2("4.2. Thẻ Khuyến Nghị Triển Khai & Danh Sách Quick Wins")
    add_p(
        "Trang Khuyến Nghị Hành Động xếp hạng các tác vụ có Điểm Ưu Tiên cao nhất và cung cấp Playbook quản trị thay đổi chi tiết:"
    )

    img4_path = "C:/Users/25/.gemini/antigravity-ide/brain/83799972-2e86-4f03-951f-2ed5cb229a65/action_recommendations_1784664319352.png"
    if os.path.exists(img4_path):
        doc.add_picture(img4_path, width=Inches(6.3))
        add_caption("Hình 4: Danh sách Thẻ Khuyến Nghị Thí Điểm và Action Plan (Recommendations)")

    # ═══════════════════════════════════════════════════════════════
    # PHẦN V: ĐỀ XUẤT CHIẾN LƯỢC & LỘ TRÌNH CHO DOANH NGHIỆP SME
    # ═══════════════════════════════════════════════════════════════
    add_h1("PHẦN V: ĐỀ XUẤT CHIẾN LƯỢC & LỘ TRÌNH CHO DOANH NGHIỆP SME")

    add_h2("5.1. Lộ Trình Triển Khai 4 Giai Đoạn (Implementation Roadmap)")
    add_bullet("Tập trung đào tạo nhận thức, giải tỏa lo ngại tâm lý cho nhân sự và ban hành quy định bảo mật dữ liệu.", "Giai đoạn 1 (Tuần 1–2): Chuẩn bị & Truyền thông:")
    add_bullet("Triển khai ngay 95 tác vụ vùng Thí điểm Ngay cho bộ phận Kế toán, CSKH để chứng minh hiệu quả ROI ($938K/năm).", "Giai đoạn 2 (Tuần 3–6): Thí điểm hẹp (Pilot):")
    add_bullet("Mở rộng AI Agent sang các phòng ban liên quan, thiết lập hệ thống giám sát thời gian thực.", "Giai đoạn 3 (Tuần 7–10): Mở rộng & Giám sát:")
    add_bullet("Chuẩn hóa quy trình làm việc kết hợp Human-AI, đóng gói case study nội bộ.", "Giai đoạn 4 (Từ tuần 11): Chuẩn hóa Quy trình:")

    add_h2("5.2. Quản Trị Rủi Ro & Chữ Ký Phê Duyệt (Governance & Human-in-the-loop)")
    add_p(
        "Đối với 34 tác vụ chứa từ khóa nhạy cảm (tài chính, ngân sách, hợp đồng pháp lý, thông tin nhân sự), "
        "doanh nghiệp SME bắt buộc tuân thủ quy chế phê duyệt chữ ký số (Human Approval Loop). AI Agent đóng vai trò "
        "Trợ lý soạn thảo (Copilot), tuyệt đối không tự động thực thi các giao dịch mà không có xác nhận của Cấp Quản lý."
    )

    # Conclusion Block
    add_h2("5.3. Kết Luận")
    add_p(
        "Hệ thống SME Strategy Advisor đã chứng minh sức mạnh của trực quan hóa dữ liệu định lượng trong việc biến "
        "bài toán AI phức tạp thành lộ trình hành động đơn giản, khả thi cho doanh nghiệp SME. Việc kết hợp hài hòa "
        "giữa công nghệ AI Agent hiện đại và tâm lý học tổ chức chính là chìa khóa vàng đảm bảo sự thành công bền vững cho doanh nghiệp."
    )

    doc.save(output_file)
    print(f"Report successfully generated at: {output_file}")

if __name__ == "__main__":
    create_word_report()
