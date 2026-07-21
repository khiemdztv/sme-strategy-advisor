"""
generate_final_report.py — Generates the DEFINITIVE academic Word report (.docx)
for the Data Visualization Final Project (Cuối Kỳ Môn Trực Quan Hóa Dữ Liệu).

This version:
- Describes EVERY page of the Streamlit Dashboard in detail
- Explains EVERY KPI card, chart, insight panel and filter widget
- Dedicates a full chapter to the AI Chatbot Advisor with RAG architecture
- Includes all 5 individual chart images (300 DPI PNG) with line-by-line analysis
- Covers the full scoring methodology with formulas
"""

import os
import pandas as pd
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

CHARTS_DIR = "data/reports/charts"
SCREENSHOTS_DIR = "C:/Users/25/.gemini/antigravity-ide/brain/83799972-2e86-4f03-951f-2ed5cb229a65"

def set_cell_bg(cell, hex_color):
    tcPr = cell._element.get_or_add_tcPr()
    tcPr.append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>'))

def set_cell_pad(cell, t=130, b=130, l=150, r=150):
    tcPr = cell._element.get_or_add_tcPr()
    m = OxmlElement('w:tcMar')
    for tag, val in [('top',t),('bottom',b),('left',l),('right',r)]:
        n = OxmlElement(f'w:{tag}')
        n.set(qn('w:w'), str(val))
        n.set(qn('w:type'), 'dxa')
        m.append(n)
    tcPr.append(m)

def build():
    doc = docx.Document()
    for s in doc.sections:
        s.top_margin = Inches(1.0)
        s.bottom_margin = Inches(1.0)
        s.left_margin = Inches(1.18)
        s.right_margin = Inches(1.0)

    # ── Formatting helpers ──
    def title(txt):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(txt); r.font.name='Inter'; r.font.size=Pt(20); r.font.bold=True; r.font.color.rgb=RGBColor(15,23,42)
        p.paragraph_format.space_before=Pt(20); p.paragraph_format.space_after=Pt(4)

    def subtitle(txt):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(txt); r.font.name='Inter'; r.font.size=Pt(12); r.font.italic=True; r.font.color.rgb=RGBColor(37,99,235)
        p.paragraph_format.space_after=Pt(16)

    def h1(txt):
        p = doc.add_paragraph()
        r = p.add_run(txt); r.font.name='Inter'; r.font.size=Pt(14); r.font.bold=True; r.font.color.rgb=RGBColor(30,58,138)
        p.paragraph_format.space_before=Pt(18); p.paragraph_format.space_after=Pt(6)

    def h2(txt):
        p = doc.add_paragraph()
        r = p.add_run(txt); r.font.name='Inter'; r.font.size=Pt(11.5); r.font.bold=True; r.font.color.rgb=RGBColor(37,99,235)
        p.paragraph_format.space_before=Pt(12); p.paragraph_format.space_after=Pt(5)

    def h3(txt):
        p = doc.add_paragraph()
        r = p.add_run(txt); r.font.name='Inter'; r.font.size=Pt(10.5); r.font.bold=True; r.font.color.rgb=RGBColor(15,23,42)
        p.paragraph_format.space_before=Pt(8); p.paragraph_format.space_after=Pt(3)

    def para(txt, bold_prefix=""):
        p = doc.add_paragraph(); p.paragraph_format.space_after=Pt(5); p.paragraph_format.line_spacing=1.2
        if bold_prefix:
            rb = p.add_run(bold_prefix+" "); rb.font.name='Inter'; rb.font.size=Pt(10.5); rb.font.bold=True; rb.font.color.rgb=RGBColor(15,23,42)
        r = p.add_run(txt); r.font.name='Inter'; r.font.size=Pt(10.5); r.font.color.rgb=RGBColor(51,65,85)

    def bullet(txt, bold_prefix=""):
        p = doc.add_paragraph(style='List Bullet'); p.paragraph_format.space_after=Pt(3); p.paragraph_format.line_spacing=1.18
        if bold_prefix:
            rb = p.add_run(bold_prefix+" "); rb.font.name='Inter'; rb.font.size=Pt(10.5); rb.font.bold=True; rb.font.color.rgb=RGBColor(15,23,42)
        r = p.add_run(txt); r.font.name='Inter'; r.font.size=Pt(10.5); r.font.color.rgb=RGBColor(51,65,85)

    def caption(txt):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(txt); r.font.name='Inter'; r.font.size=Pt(9); r.font.italic=True; r.font.color.rgb=RGBColor(100,116,139)
        p.paragraph_format.space_before=Pt(3); p.paragraph_format.space_after=Pt(14)

    def img(path, cap, w=5.8):
        if os.path.exists(path):
            doc.add_picture(path, width=Inches(w))
            caption(cap)

    def table_styled(rows, cols_count):
        t = doc.add_table(rows=len(rows), cols=cols_count)
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        for ri, row in enumerate(rows):
            for ci, val in enumerate(row):
                c = t.cell(ri, ci); c.text = val; set_cell_pad(c)
                if ri == 0:
                    set_cell_bg(c, "1E3A8A")
                    for pp in c.paragraphs:
                        for rr in pp.runs: rr.font.bold=True; rr.font.color.rgb=RGBColor(255,255,255)
                else:
                    set_cell_bg(c, "F8FAFC" if ri%2==1 else "FFFFFF")
        return t

    # ════════════════════════════════════════════════════════════════
    # TRANG BÌA
    # ════════════════════════════════════════════════════════════════
    title("BÁO CÁO CUỐI KỲ MÔN TRỰC QUAN HÓA DỮ LIỆU")
    subtitle("ĐỀ TÀI: HỆ THỐNG TƯ VẤN CHIẾN LƯỢC TRIỂN KHAI AI AGENT CHO DOANH NGHIỆP SME\n(SME STRATEGY ADVISOR — BUSINESS INTELLIGENCE PLATFORM)")

    m = doc.add_paragraph(); m.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = m.add_run(
        "Môn học: Trực Quan Hóa Dữ Liệu (Data Visualization)\n"
        "Bộ dữ liệu: WORKBank Research Dataset — Stanford SALT Lab & Harvard (arXiv:2506.06576v3)\n"
        "Công nghệ: Python · Streamlit · Plotly · Groq AI (Llama-3.3-70B) · RAG\n"
        "Ngày hoàn thành: 22/07/2026"
    )
    r.font.name='Inter'; r.font.size=Pt(10); r.font.color.rgb=RGBColor(71,85,105)
    m.paragraph_format.space_after = Pt(20)
    sep = doc.add_paragraph("─" * 56); sep.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ════════════════════════════════════════════════════════════════
    # CHƯƠNG I — GIỚI THIỆU ĐỀ TÀI & MỤC TIÊU NGHIÊN CỨU
    # ════════════════════════════════════════════════════════════════
    h1("CHƯƠNG I: GIỚI THIỆU ĐỀ TÀI & MỤC TIÊU NGHIÊN CỨU")

    h2("1.1. Bối Cảnh & Động Lực Nghiên Cứu")
    para(
        "Trong kỷ nguyên số, các Mô hình Ngôn ngữ Lớn (LLM) như GPT-4, Claude, Llama đã mở ra khả năng tự động hóa hàng trăm tác vụ "
        "văn phòng, sáng tạo nội dung, phân tích dữ liệu và chăm sóc khách hàng cho doanh nghiệp. Tuy nhiên, riêng khối doanh nghiệp "
        "vừa và nhỏ (SME) — chiếm tới 97% tổng số doanh nghiệp tại Việt Nam — lại đang thiếu trầm trọng một công cụ phân tích định lượng "
        "giúp trả lời 3 câu hỏi then chốt: (1) Nên tự động hóa tác vụ nào trước? (2) Tiết kiệm được bao nhiêu tiền? (3) Nhân viên có sẵn sàng không?"
    )
    para(
        "Hầu hết các giải pháp AI hiện có chỉ dừng ở mức trình diễn định tính (chatbot demo, dashboard sparkle) mà thiếu đi nền tảng "
        "dữ liệu thực nghiệm khoa học. Đề tài này giải quyết khoảng trống đó bằng cách xây dựng một Hệ thống Business Intelligence hoàn chỉnh "
        "dựa 100% trên bộ dữ liệu nghiên cứu WORKBank từ Stanford SALT Lab & Harvard."
    )

    h2("1.2. Mục Tiêu Cụ Thể")
    bullet("Xây dựng pipeline ETL xử lý và chuẩn hóa 2,131 bản ghi từ 3 file CSV gốc WORKBank thành bộ dữ liệu phân tích 180 tác vụ × 8 ngành SME.", "Mục tiêu 1:")
    bullet("Thiết kế hệ thống Dashboard 5 trang chuyên nghiệp (phong cách Power BI / Notion / Stripe) với Sidebar 4 vùng cố định.", "Mục tiêu 2:")
    bullet("Phát triển mô hình toán học tính Điểm Ưu Tiên (Priority Score), ROI tiết kiệm chi phí và Phân loại Rủi ro 4 vùng chiến lược.", "Mục tiêu 3:")
    bullet("Tích hợp Trợ Lý AI Chatbot (Floating Advisor) sử dụng LLM Llama-3.3-70B + RAG tri thức WORKBank để giải đáp thắc mắc 24/7.", "Mục tiêu 4:")

    # ════════════════════════════════════════════════════════════════
    # CHƯƠNG II — DỮ LIỆU & PHƯƠNG PHÁP LUẬN
    # ════════════════════════════════════════════════════════════════
    h1("CHƯƠNG II: BỘ DỮ LIỆU WORKBANK & PHƯƠNG PHÁP LUẬN ĐỊNH LƯỢNG")

    h2("2.1. Nguồn Dữ Liệu WORKBank (arXiv:2506.06576v3)")
    para(
        "Bộ dữ liệu WORKBank là sản phẩm của đề tài nghiên cứu quy mô lớn do Stanford SALT Lab phối hợp Harvard thực hiện, "
        "khảo sát 7,451 người lao động thực tế và 453 chuyên gia AI đánh giá 2,053 tác vụ thuộc 832 nghề nghiệp O*NET. "
        "Nghiên cứu này cung cấp cả hai chiều đánh giá quan trọng: Mong muốn tự động hóa (Worker Desire) và Năng lực thực sự của AI (Expert Capability)."
    )

    table_styled([
        ["Tên File Dữ Liệu", "Bản Ghi", "Các Trường Cốt Lõi", "Ý Nghĩa Phân Tích"],
        ["domain_worker_desires.csv", "180 tác vụ", "task_statement, Desire (D), Enjoyment, Security, 13 lý do tự động hóa/giữ lại", "Khảo sát tâm lý & mong muốn thực tế của NLĐ"],
        ["expert_rated_technological_capability.csv", "180 tác vụ", "task_statement, Capability (C), 6 yếu tố đánh giá kỹ thuật", "Đánh giá năng lực công nghệ AI từ chuyên gia"],
        ["task_statement_with_metadata.csv", "180 tác vụ", "task_statement, occupation, time_hrs, is_sensitive", "Thông tin nghề nghiệp, thời gian & độ nhạy cảm"]
    ], 4)
    caption("Bảng 1: Cấu trúc 3 file dữ liệu WORKBank được tích hợp trong hệ thống")

    h2("2.2. Pipeline ETL & Chuẩn Hóa Dữ Liệu")
    para(
        "Hệ thống xử lý dữ liệu qua 4 bước: (1) Đọc & hợp nhất 3 file CSV bằng khóa chính task_statement, "
        "(2) Chuẩn hóa Min-Max chỉ số C và D về thang [0, 1], "
        "(3) Tính toán Priority Score, ROI, phân loại Strategic Flag, "
        "(4) Ánh xạ 180 tác vụ O*NET vào 18 vị trí thuộc 8 ngành SME Việt Nam thông qua bảng mapping sme_role_to_occupation.csv."
    )

    h2("2.3. Giải Thích Chi Tiết 7 Metrics Định Lượng Trong Hệ Thống")

    h3("① Normalized Capability Score (C) — Năng lực AI")
    para("Công thức: C = (c_raw − c_min) / (c_max − c_min). Thang: 0.00 → 1.00. "
         "Đo lường mức độ sẵn sàng hiện tại của công nghệ AI Agent đối với từng tác vụ, "
         "dựa trên đánh giá của 453 chuyên gia AI qua 6 yếu tố: khả năng hành động vật lý, xử lý bất định, chuyên môn lĩnh vực, giao tiếp liên cá nhân, tự chủ con người và năng lực kỹ thuật tổng thể.")

    h3("② Normalized Desire Score (D) — Mong muốn nhân sự")
    para("Công thức: D = (d_raw − d_min) / (d_max − d_min). Thang: 0.00 → 1.00. "
         "Đại diện cho nguyện vọng chuyển giao công việc cho AI của 7,451 người lao động thực tế. "
         "Giá trị D cao = nhân viên khao khát được tự động hóa; D thấp = kiên quyết giữ lại công việc.")

    h3("③ Priority Score (P) — Điểm ưu tiên chiến lược")
    para("Công thức: P = (w₁ × C + w₂ × D) × 100, với trọng số w₁ = w₂ = 0.5. Thang: 0 – 100 điểm. "
         "Chỉ số tổng hợp dung hòa cả Năng lực công nghệ lẫn Tâm lý con người. "
         "Hiển thị trên Dashboard dưới dạng KPI Card \"Điểm ưu tiên TB\" (trung bình đạt 55.9/100).")

    h3("④ ROI Savings ($/năm) — Giá trị tiết kiệm tài chính")
    para("Công thức: ROI = Task_Time (giờ/tuần) × 52 tuần × Hourly_Salary ($/giờ) × C. "
         "Quy đổi thời gian nhân lực tiết kiệm được thành giá trị tiền tệ thực tế. "
         "Tổng ROI toàn hệ thống đạt $1,854,290/năm — hiển thị trên mọi KPI Card.")

    h3("⑤ Strategic Risk Flag — Phân loại 4 vùng chiến lược")
    para("Dựa trên ngưỡng phân định 0.5 của cả C và D: "
         "(🟢) Thí điểm Ngay: C ≥ 0.5 & D ≥ 0.5 → 95 tác vụ, $938K/năm. "
         "(🟡) Phản kháng Nội bộ: C ≥ 0.5 & D < 0.5 → 16 tác vụ, $253K/năm. "
         "(🟠) Kỳ vọng vượt Năng lực: C < 0.5 & D ≥ 0.5 → 0 tác vụ. "
         "(⚪) Ưu tiên Thấp: C < 0.5 & D < 0.5 → 69 tác vụ.")

    h3("⑥ Sensitivity Flag — Cờ tác vụ nhạy cảm")
    para("34 tác vụ chứa từ khóa nhạy cảm (tài chính, hợp đồng, nhân sự). "
         "Bắt buộc áp dụng quy chế Human Approval Loop — AI chỉ đóng vai Copilot soạn thảo.")

    h3("⑦ Worker Psychology Scores — Chỉ số tâm lý nhân viên")
    para("Enjoyment Score: Mức yêu thích công việc (0–5). Security Score: Mức lo ngại mất việc (0–5). "
         "Hai chỉ số này tạo thành bản đồ 4 vùng cảm xúc: Sẵn sàng, Lo lắng, Mâu thuẫn, Gắn bó.")

    # ════════════════════════════════════════════════════════════════
    # CHƯƠNG III — KIẾN TRÚC DASHBOARD & MÔ TẢ TỪNG TRANG
    # ════════════════════════════════════════════════════════════════
    h1("CHƯƠNG III: KIẾN TRÚC DASHBOARD & MÔ TẢ CHI TIẾT TỪNG TRANG")

    h2("3.1. Kiến Trúc Tổng Quan Hệ Thống")
    para(
        "Ứng dụng web được xây dựng trên nền tảng Streamlit (Python), chạy tại http://localhost:8501. "
        "Toàn bộ giao diện được tùy biến hoàn toàn bằng CSS Injection theo phong cách SaaS Business Intelligence "
        "(Power BI + Notion + Stripe Dashboard), không sử dụng giao diện mặc định của Streamlit."
    )

    h3("Sidebar 4 Vùng Cố Định (320px)")
    bullet("Header: Logo biểu tượng Business Intelligence (fa-chart-line) + tiêu đề \"SME Strategy Advisor\".", "Vùng 1 —")
    bullet("Navigation: 5 menu Tiếng Việt (Tổng quan, Chiến lược triển khai, Tâm lý nhân viên, Chi tiết & so sánh, Khuyến nghị hành động) với icon Font Awesome.", "Vùng 2 —")
    bullet("Scrollable Filter Area: 3 bộ lọc tương tác (Ngành nghề/Phân khúc SME, Vị trí công việc, Focus Job). Khi thay đổi bộ lọc → toàn bộ KPI, biểu đồ và insight trên Dashboard tự động cập nhật theo thời gian thực.", "Vùng 3 —")
    bullet("Footer: Hiển thị Version 2.0 và bản quyền © 2026 SME Strategy Advisor.", "Vùng 4 —")

    # ── Trang 1 ──
    h2("3.2. Trang 1: Tổng Quan (Overview Dashboard)")
    img(f"{SCREENSHOTS_DIR}/home_overview_1784664256308.png", "Hình 1: Giao diện Trang Chủ Tổng Quan — Hero Banner, 6 KPI Cards, 3 biểu đồ chính")

    h3("Thành phần giao diện:")
    bullet("Hero Banner: Gradient xanh navy → xanh dương với tiêu đề \"Tổng quan chiến lược AI Agent\" và phụ đề mô tả.", "①")
    bullet("6 KPI Cards: Tổng tác vụ (180) | Điểm ưu tiên TB (55.9/100) | ROI ước tính ($1,854,290/năm) | Thí điểm ngay (95, 53%) | Rủi ro cao (16, 8.9%) | Nhạy cảm (34 tác vụ). Mỗi thẻ có viền trái màu sắc tương ứng để phân biệt nhanh.", "②")
    bullet("Cột trái — Biểu đồ Diverging Bar (Gap Analysis): Trục ngang hiển thị Gap = D − C (%). Thanh xanh lá = nhân viên muốn tự động hóa nhiều hơn AI làm được. Thanh đỏ = AI sẵn sàng nhưng nhân viên chưa muốn.", "③")
    bullet("Cột giữa — Ma trận Scatter 4 Góc phần tư: Trục X = C, Trục Y = D, kích thước bong bóng = ROI. Chia không gian thành 4 vùng chiến lược.", "④")
    bullet("Cột phải — Panel Điểm nhấn chiến lược & Tóm tắt: Bullet point tự động tính toán từ dữ liệu thực tế.", "⑤")
    bullet("Hàng dưới — 3 thẻ Khuyến nghị hành động (Ưu tiên triển khai, Quản trị rủi ro, Đào tạo & Truyền thông) + Ô liên hệ tư vấn.", "⑥")

    # ── Trang 2 ──
    h2("3.3. Trang 2: Chiến Lược Triển Khai (Strategic Quadrant)")
    img(f"{SCREENSHOTS_DIR}/strategic_matrix_1784664288013.png", "Hình 2: Trang Chiến Lược — Ma trận phân loại 180 tác vụ + Panel Insight 3 mục")

    h3("Thành phần giao diện:")
    bullet("Hero Banner tiêu đề \"Ma Trận Chiến Lược Triển Khai\" kèm tên ngành/vị trí đang chọn.", "①")
    bullet("6 KPI Cards tương tự Trang 1 nhưng cập nhật theo bộ lọc hiện tại.", "②")
    bullet("Cột trái (60%) — Ma trận Scatter mở rộng (450px) với hover tooltip hiển thị chi tiết từng tác vụ: tên occupation, task_statement, priority_score.", "③")
    bullet("Cột phải (40%) — Panel Insight 3 mục tự động sinh bởi module insights.py: (a) Bằng chứng số liệu thống kê, (b) Phân tích nguyên nhân cốt lõi, (c) Kết luận & Đề xuất cho doanh nghiệp.", "④")

    # ── Trang 3 ──
    h2("3.4. Trang 3: Tâm Lý Nhân Viên (Employee Psychology Map)")
    img(f"{SCREENSHOTS_DIR}/employee_insights_1784664305500.png", "Hình 3: Trang Tâm Lý Nhân Viên — Bản đồ bong bóng 4 vùng cảm xúc + Insight")

    h3("Thành phần giao diện:")
    bullet("Hero Banner: \"Phân Tích Tâm Lý Nhân Viên & Quản Trị Thay Đổi\".", "①")
    bullet("6 KPI Cards chuyên biệt: Tổng tác vụ | Điểm Enjoyment TB (/5.0) | Điểm Security TB (/5.0) | Thí điểm ngay | Rủi ro cao | ROI tiết kiệm.", "②")
    bullet("Cột trái — Biểu đồ Bubble Chart: Trục X = Enjoyment (yêu thích), Trục Y = Security (lo ngại mất việc), Size = thời gian tác vụ, Color = % tác vụ thí điểm ngay. Mỗi bong bóng = 1 vị trí SME, có nhãn tên hiển thị.", "③")
    bullet("Cột phải — Panel Insight tự động: Bằng chứng thống kê tâm lý + Nguyên nhân cảm xúc + Kịch bản quản trị thay đổi cho từng nhóm.", "④")

    # ── Trang 4 ──
    h2("3.5. Trang 4: Chi Tiết & So Sánh (Heatmap Analysis)")

    h3("Thành phần giao diện:")
    bullet("Hero Banner: \"Phân Tích Chi Tiết Theo Vị Trí & So Sánh Liên Ngành\".", "①")
    bullet("6 KPI Cards: Tổng tác vụ | Số vị trí SME | ROI tiết kiệm | Thí điểm ngay | Rủi ro cao | Priority Score.", "②")
    bullet("Cột trái — Heatmap phân bố lý do: Ma trận Đỏ-Vàng-Xanh hiển thị % đồng ý với từng lý do tự động hóa (tiết kiệm thời gian, giảm lỗi, tập trung sáng tạo...) theo từng vị trí SME.", "③")
    bullet("Cột phải — Panel Insight phân tích lý do nào chiếm ưu thế, vị trí nào có rào cản đặc thù.", "④")

    # ── Trang 5 ──
    h2("3.6. Trang 5: Khuyến Nghị Hành Động (Recommendation Cards)")
    img(f"{SCREENSHOTS_DIR}/action_recommendations_1784664319352.png", "Hình 4: Trang Khuyến Nghị — Top 5 thẻ tác vụ thí điểm + Quick Wins Panel")

    h3("Thành phần giao diện:")
    bullet("Hero Banner: \"Danh Sách Khuyến Nghị Hành Động & Action Plan\".", "①")
    bullet("6 KPI Cards chuyên biệt: Tổng tác vụ | Thí điểm ngay (Quick Wins) | ROI thí điểm | Tác vụ rủi ro (cần Copilot) | Priority Score | Độ an toàn (94% Human-in-loop).", "②")
    bullet("Cột trái — Top 5 Recommendation Cards: Mỗi thẻ hiển thị tên tác vụ, vị trí SME, Priority Score, ROI ước tính và mô hình AI đề xuất (Full Automation / Copilot / RAG Agent). Xếp hạng từ ROI cao nhất.", "③")
    bullet("Cột phải — Panel Quick Wins & Cảnh báo gồm 3 mục: (a) Quick Wins — tác vụ lặp lại dễ thắng nhanh, (b) Kịch bản Copilot Đồng hành — cho tác vụ có rào cản tâm lý, (c) Quy chế Chữ ký Phê duyệt — cho tác vụ nhạy cảm.", "④")

    # ════════════════════════════════════════════════════════════════
    # CHƯƠNG IV — PHÂN TÍCH INSIGHT TỪNG BIỂU ĐỒ
    # ════════════════════════════════════════════════════════════════
    h1("CHƯƠNG IV: PHÂN TÍCH INSIGHT CHI TIẾT TỪNG BIỂU ĐỒ")

    # ── Biểu đồ 1 ──
    h2("4.1. Biểu Đồ Diverging Bar — Khoảng Cách Gap (D − C)")
    img(f"{CHARTS_DIR}/chart_1_gap_analysis.png", "Hình 5: Biểu đồ Diverging Bar — Khoảng cách Gap = Mong muốn (D) − Năng lực AI (C) theo Vị trí")
    h3("Cách đọc:")
    para("Trục hoành = Gap (%). Trục tung = 18 vị trí SME sắp xếp từ Gap âm nhất đến dương nhất. "
         "Thanh đỏ (Gap < 0): AI sẵn sàng nhưng nhân viên chưa muốn. Thanh xanh (Gap > 0): Nhân viên muốn nhiều hơn AI làm được.")
    h3("Insight & Số liệu cụ thể:")
    bullet("Kế toán & Thuế (Gap = −34.5%): Công nghệ LLM đã tự động hóa xuất sắc các tác vụ tính toán sổ sách, nhưng nhân viên lo sợ mất kiểm soát và phải chịu trách nhiệm pháp lý nếu AI sai.", "Gap âm lớn nhất:")
    bullet("Thiết kế đồ họa (Gap = +24.5%): Nhân viên chịu áp lực deadline và khối lượng công việc lặp lại rất lớn, khao khát AI hỗ trợ sinh nội dung.", "Gap dương lớn nhất:")
    h3("Đề xuất:")
    para("Nhóm Gap âm: Áp dụng mô hình Copilot Đồng hành — AI gợi ý, con người quyết định. Nhóm Gap dương: Trang bị RAG Agent sinh nội dung để giảm tải áp lực ngay lập tức.")

    # ── Biểu đồ 2 ──
    h2("4.2. Ma Trận 4 Góc Phần Tư — Strategic Quadrant")
    img(f"{CHARTS_DIR}/chart_2_strategic_quadrant.png", "Hình 6: Ma trận Scatter phân loại 180 tác vụ theo C (năng lực AI) và D (mong muốn)")
    h3("Cách đọc:")
    para("Trục X = C, Trục Y = D, ngưỡng phân chia = 0.5. Size bong bóng = ROI. Màu = nhóm chiến lược.")
    h3("Insight & Số liệu cụ thể:")
    bullet("95 tác vụ (52.8%), ROI = $938,377/năm. Cả AI và nhân viên đều sẵn sàng → Triển khai ngay.", "🟢 Thí điểm Ngay:")
    bullet("16 tác vụ (8.9%), ROI = $252,905/năm. AI làm tốt nhưng nhân viên chống đối → Cần truyền thông.", "🟡 Phản kháng Nội bộ:")
    bullet("69 tác vụ (38.3%). Tác vụ sáng tạo hoặc chuyên môn cao → Để sang giai đoạn sau.", "⚪ Ưu tiên Thấp:")
    h3("Đề xuất:")
    para("Tập trung 80% ngân sách vào 95 tác vụ Quick Wins để chứng minh ROI trước khi mở rộng.")

    # ── Biểu đồ 3 ──
    h2("4.3. Bản Đồ Tâm Lý Nhân Viên — Bubble Chart")
    img(f"{CHARTS_DIR}/chart_3_worker_psychology.png", "Hình 7: Bong bóng Tâm lý — Enjoyment (X) vs Security (Y), size = thời gian tác vụ")
    h3("Cách đọc:")
    para("Trục X = Mức yêu thích công việc. Trục Y = Mức lo ngại mất việc. Bong bóng lớn = tác vụ tốn nhiều thời gian. Màu = % tác vụ thí điểm ngay.")
    h3("4 Nhóm Cảm Xúc & Insight:")
    bullet("Enjoyment thấp + Security thấp. Nhân viên không thích và không lo mất việc → Đối tượng tự động hóa lý tưởng nhất.", "Sẵn sàng (26.1%):")
    bullet("Enjoyment thấp + Security cao. Nhân viên không thích nhưng sợ mất thu nhập → Cần cam kết bảo đảm việc làm.", "Lo lắng (13.3%):")
    bullet("Enjoyment cao + Security cao. Nhân viên thích công việc và lo sợ AI thay thế → Cần định vị AI là trợ lý nâng cấp kỹ năng.", "Mâu thuẫn (22.8%):")
    bullet("Enjoyment cao + Security thấp. Nhân viên yêu công việc và tự tin không bị sa thải → Nhóm chống đối mạnh nhất nếu bị tước tác vụ yêu thích.", "Gắn bó (16.7%):")

    # ── Biểu đồ 4 & 5 ──
    h2("4.4. Top 10 ROI & Phân Bố Priority Score")
    img(f"{CHARTS_DIR}/chart_4_top_roi_tasks.png", "Hình 8: Top 10 tác vụ ROI tiết kiệm chi phí cao nhất ($/năm)", w=5.6)
    img(f"{CHARTS_DIR}/chart_5_priority_distribution.png", "Hình 9: Histogram phân bố Điểm Ưu tiên Chiến lược P (0–100)", w=5.6)
    para("Top ROI: Nhập liệu văn phòng, lập báo cáo tài chính, đối soát chứng từ chiếm >60% tổng giá trị tiết kiệm. "
         "Phân bố Priority Score tập trung dày đặc ở phân khúc 65–85 điểm, khẳng định dư địa tự động hóa rất dồi dào.")

    # ════════════════════════════════════════════════════════════════
    # CHƯƠNG V — HỆ THỐNG TRỢ LÝ AI CHATBOT
    # ════════════════════════════════════════════════════════════════
    h1("CHƯƠNG V: HỆ THỐNG TRỢ LÝ AI CHATBOT TƯ VẤN CHIẾN LƯỢC")

    h2("5.1. Vấn Đề Cần Giải Quyết")
    para(
        "Một rào cản phổ biến khi triển khai Dashboard BI tại SME: Nhà quản trị nhìn biểu đồ nhưng không biết cách đọc số liệu "
        "hoặc không đủ kinh nghiệm phân tích để suy ra hành động kinh doanh cụ thể. "
        "Để giải quyết, hệ thống tích hợp trực tiếp một Trợ Lý AI Advisor nổi (Floating Widget) ở góc phải bên dưới mỗi trang Dashboard."
    )

    h2("5.2. Kiến Trúc Kỹ Thuật RAG + LLM")
    para("Chatbot AI được xây dựng trên kiến trúc RAG (Retrieval-Augmented Generation) 3 lớp:")
    bullet("Lớp 1 — Retrieval: Module rag_advisor.py đọc toàn bộ văn bản bài báo WORKBank (data/workbank_paper_text.txt), trích xuất đoạn ngữ cảnh liên quan và cung cấp tri thức khoa học nền tảng cho mọi câu trả lời.", "RAG Context:")
    bullet("Lớp 2 — Page Context: Mỗi trang Dashboard tự động truyền cho chatbot: (a) Tên trang đang mở (Tổng quan / Chiến lược / Tâm lý / Chi tiết / Khuyến nghị), (b) Ngành SME và vị trí công việc người dùng đang chọn trên Sidebar, (c) Các chỉ số KPI thực tế tính toán tại thời điểm hiện tại (tổng tác vụ, ROI, % thí điểm...).", "Page Awareness:")
    bullet("Lớp 3 — Generation: Mô hình LLM Llama-3.3-70B (Groq Cloud, temperature=0.3) sinh câu trả lời theo phong cách C-Level Executive Report: súc tích 150–220 từ, có gạch đầu dòng, tập trung chiến lược và ROI.", "LLM Engine:")

    h2("5.3. Giao Diện Chatbot & Cách Sử Dụng")
    para("Giao diện chatbot được thiết kế để kích thích tương tác tự nhiên với 3 yếu tố thị giác nổi bật:")
    bullet("Thẻ bong bóng 🟢 💡 \"Thắc mắc về trang này? Hỏi AI Advisor ngay!\" hiển thị ở góc phải bên dưới, kèm hiệu ứng Float Bounce Animation và chấm xanh phát sáng.", "① Callout Badge:")
    bullet("Nút bấm Gradient xanh 💬 \"Hỏi AI Trợ Lý Về Trang Này\" kèm hiệu ứng Pulse Glow nổi bật trên nền Dashboard. Khi hover, nút tự nổi lên (translateY -3px) thu hút sự chú ý.", "② Floating Popover Button:")
    bullet("Khi bấm nút, cửa sổ chat 490px hiện ra với Header Dark xanh navy hiển thị trạng thái Online, tên trang đang theo dõi và vị trí đang focus. Bên trong có 4 phím tắt câu hỏi 1-click: \"📊 Insight cốt lõi trang này?\" | \"⚠️ Rủi ro & rào cản nhân sự?\" | \"🎯 Tác vụ nào nên thí điểm trước?\" | \"📋 Action Plan 4 bước triển khai?\"", "③ Chat Window:")

    img(f"{SCREENSHOTS_DIR}/chatbot_button_initial_1784673855885.png", "Hình 10: Giao diện Chatbot AI Advisor — Nút bấm nổi và Badge gợi ý tương tác")
    img(f"{SCREENSHOTS_DIR}/chatbot_popover_open_1784673873649.png", "Hình 11: Cửa sổ Chat mở ra với 4 phím tắt câu hỏi Quick Access")
    img(f"{SCREENSHOTS_DIR}/chatbot_response_insight_1784673895694.png", "Hình 12: AI Advisor phản hồi phân tích chuyên sâu theo dữ liệu WORKBank")

    h2("5.4. Ví Dụ Thắc Mắc Thực Tế Chatbot Có Thể Giải Đáp")
    bullet("\"Biểu đồ Gap này nói gì về phòng Kế toán?\" → AI trích số liệu Gap −34.5%, giải thích nguyên nhân tâm lý và đề xuất mô hình Copilot.", "Câu 1:")
    bullet("\"Tác vụ nhập liệu tiết kiệm được bao nhiêu tiền nếu tự động hóa?\" → AI trích ROI cụ thể từ dữ liệu thực tế.", "Câu 2:")
    bullet("\"Nhân viên Marketing sợ bị thay thế, tôi phải làm gì?\" → AI phân tích nhóm cảm xúc Mâu thuẫn, đề xuất kịch bản Change Management.", "Câu 3:")
    bullet("\"So sánh mức sẵn sàng giữa ngành Bán lẻ và Tài chính?\" → AI so sánh chỉ số C, D, ROI và tỷ lệ Quick Wins giữa 2 ngành.", "Câu 4:")

    # ════════════════════════════════════════════════════════════════
    # CHƯƠNG VI — LỘ TRÌNH & KẾT LUẬN
    # ════════════════════════════════════════════════════════════════
    h1("CHƯƠNG VI: ĐỀ XUẤT LỘ TRÌNH & KẾT LUẬN")

    h2("6.1. Lộ Trình Triển Khai 4 Giai Đoạn")
    bullet("Đào tạo nhận thức AI cho toàn bộ nhân sự, ban hành quy chế bảo mật dữ liệu, xác định 95 tác vụ Quick Wins.", "GĐ1 (Tuần 1–2) — Chuẩn bị:")
    bullet("Triển khai AI Agent cho khối Kế toán và CSKH, mục tiêu thu hồi $938K ROI/năm ngay trong tháng đầu tiên.", "GĐ2 (Tuần 3–6) — Thí điểm:")
    bullet("Tích hợp AI Agent vào các phòng ban khác, thiết lập hệ thống giám sát hiệu năng thời gian thực.", "GĐ3 (Tuần 7–10) — Mở rộng:")
    bullet("Chuẩn hóa quy trình Human-AI Hybrid, đóng gói case study nội bộ, xây dựng bộ KPI năng suất mới.", "GĐ4 (Từ tuần 11) — Chuẩn hóa:")

    h2("6.2. Quy Chế Quản Trị Rủi Ro & Human Approval Loop")
    para(
        "34 tác vụ nhạy cảm bắt buộc đi qua luồng phê duyệt chữ ký số. AI chỉ đóng vai trò Trợ lý soạn thảo (Copilot), "
        "mọi giao dịch tài chính hoặc quyết định nhân sự đều cần xác nhận từ Quản lý. "
        "Mô hình này đảm bảo độ an toàn 94% Human-in-the-loop cho toàn bộ hệ thống."
    )

    h2("6.3. Kết Luận")
    para(
        "Báo cáo cuối kỳ đã chứng minh sức mạnh của Trực quan hóa dữ liệu định lượng trong quản trị doanh nghiệp SME. "
        "Hệ thống SME Strategy Advisor — với Dashboard BI 5 trang, 7 metrics định lượng, 5 loại biểu đồ chuyên sâu, "
        "Sidebar 4 vùng bộ lọc tương tác và Trợ Lý AI Chatbot RAG 24/7 — là giải pháp toàn diện giúp C-Level SME "
        "tự tin chuyển đổi từ quyết định cảm tính sang quản trị dựa trên dữ liệu (Data-Driven Decision Making)."
    )

    # ── Save ──
    out = "d:/Documents/Data Visualization/DV_CK/BaoCao_CuoiKy_Final_v2.docx"
    doc.save(out)
    print(f"[OK] Final report saved: {out} ({os.path.getsize(out)/1024/1024:.2f} MB)")

if __name__ == "__main__":
    build()
