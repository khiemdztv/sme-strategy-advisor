"""
generate_comprehensive_word_report.py — Generates an extensive academic Word report (.docx)
for the Data Visualization Final Project.
Includes detailed metrics explanations, chart insights, individual chart figures, and a dedicated chapter on the Floating AI Chatbot Strategy Advisor.
"""

import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=140, bottom=140, left=180, right=180):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def create_comprehensive_report(output_file="d:/Documents/Data Visualization/DV_CK/BaoCao_HoanChinh_CuoiKy_TrucQuanHoaDuLieu_AI_Agent_SME.docx"):
    doc = docx.Document()
    
    # ── Page Margins (Standard Academic 2.54 cm = 1 inch) ──
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # ── Style Helpers ──
    def add_title(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.name = 'Inter'
        run.font.size = Pt(21)
        run.font.bold = True
        run.font.color.rgb = RGBColor(15, 23, 42)
        p.paragraph_format.space_before = Pt(24)
        p.paragraph_format.space_after = Pt(6)

    def add_subtitle(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.name = 'Inter'
        run.font.size = Pt(12.5)
        run.font.italic = True
        run.font.color.rgb = RGBColor(37, 99, 235)
        p.paragraph_format.space_after = Pt(20)

    def add_h1(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = 'Inter'
        run.font.size = Pt(14.5)
        run.font.bold = True
        run.font.color.rgb = RGBColor(30, 58, 138)
        p.paragraph_format.space_before = Pt(20)
        p.paragraph_format.space_after = Pt(8)

    def add_h2(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = 'Inter'
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = RGBColor(37, 99, 235)
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)

    def add_h3(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = 'Inter'
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = RGBColor(15, 23, 42)
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)

    def add_p(text, bold_prefix=""):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.2
        if bold_prefix:
            run_b = p.add_run(bold_prefix + " ")
            run_b.font.name = 'Inter'
            run_b.font.size = Pt(10.5)
            run_b.font.bold = True
            run_b.font.color.rgb = RGBColor(15, 23, 42)
        run = p.add_run(text)
        run.font.name = 'Inter'
        run.font.size = Pt(10.5)
        run.font.color.rgb = RGBColor(51, 65, 85)
        return p

    def add_bullet(text, bold_prefix=""):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.18
        if bold_prefix:
            run_b = p.add_run(bold_prefix + " ")
            run_b.font.name = 'Inter'
            run_b.font.size = Pt(10.5)
            run_b.font.bold = True
            run_b.font.color.rgb = RGBColor(15, 23, 42)
        run = p.add_run(text)
        run.font.name = 'Inter'
        run.font.size = Pt(10.5)
        run.font.color.rgb = RGBColor(51, 65, 85)

    def add_caption(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.name = 'Inter'
        run.font.size = Pt(9.5)
        run.font.italic = True
        run.font.color.rgb = RGBColor(100, 116, 139)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(16)

    def add_image_if_exists(img_path, caption_text, width_inches=5.9):
        if os.path.exists(img_path):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(10)
            doc.add_picture(img_path, width=Inches(width_inches))
            add_caption(caption_text)

    # ═══════════════════════════════════════════════════════════════
    # BÌA VÀ THÔNG TIN CHUNG
    # ═══════════════════════════════════════════════════════════════
    add_title("BÁO CÁO NGHIÊN CỨU TRỰC QUAN HÓA DỮ LIỆU CUỐI KỲ HOÀN CHỈNH")
    add_subtitle("HỆ THỐNG TƯ VẤN ĐỊNH HƯỚNG & CHIẾN LƯỢC TRIỂN KHAI AI AGENT CHO DOANH NGHIỆP SME\n(SME STRATEGY ADVISOR — BUSINESS INTELLIGENCE PLATFORM)")

    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_meta = meta_p.add_run(
        "Môn học: Trực quan hóa dữ liệu (Data Visualization)\n"
        "Nguồn dữ liệu chuẩn mực: WORKBank Research Dataset (Stanford SALT Lab & Harvard, arXiv:2506.06576v3)\n"
        "Đơn vị thực hiện: Nhóm Nghiên Cứu Trực Quan Hóa Dữ Liệu SME Advisor\n"
        "Ngày hoàn thành: 22/07/2026"
    )
    r_meta.font.name = 'Inter'
    r_meta.font.size = Pt(10.5)
    r_meta.font.color.rgb = RGBColor(71, 85, 105)
    meta_p.paragraph_format.space_after = Pt(24)

    doc.add_paragraph("─" * 58).alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ═══════════════════════════════════════════════════════════════
    # CHƯƠNG I: GIỚI THIỆU TỔNG QUAN & TÍNH CẤP THIẾT CỦA ĐỀ TÀI
    # ═══════════════════════════════════════════════════════════════
    add_h1("CHƯƠNG I: GIỚI THIỆU TỔNG QUAN & TÍNH CẤP THIẾT CỦA ĐỀ TÀI")
    
    add_h2("1.1. Bối Cảnh Thực Tiễn Tự Động Hóa Trong Doanh Nghiệp SME")
    add_p(
        "Làn sóng trí tuệ nhân tạo thế hệ mới với sự bùng nổ của các Mô hình Ngôn ngữ Lớn (LLM) và hệ thống AI Agent đa tác vụ "
        "đang tạo ra những cơ hội tái cấu trúc năng suất chưa từng có cho khối doanh nghiệp vừa và nhỏ (SME). "
        "Không giống các tập đoàn đa quốc gia sở hữu nguồn lực tài chính và đội ngũ chuyên gia công nghệ hùng hậu, "
        "các doanh nghiệp SME phải đối mặt với áp lực duy trì dòng tiền, bộ máy nhân sự tinh gọn và tối ưu hóa chi phí vận hành từng ngày."
    )
    add_p(
        "Tuy nhiên, quá trình triển khai AI Agent tại các SME đang vấp phải những rào cản nghiêm trọng: "
        "Đa số cấp quản lý ra quyết định dựa trên cảm tính hoặc sự thôi thúc tạm thời từ các công cụ trình diễn (AI Magic/Sparkles), "
        "mà thiếu đi một lộ trình định lượng rõ ràng về giá trị tiết kiệm chi phí (ROI), rào cản tâm lý người lao động và rủi ro an toàn thông tin."
    )

    add_h2("1.2. Tính Cấp Thiết & Mục Tiêu Nghiên Cứu Khoa Học")
    add_p(
        "Đề tài hướng tới việc giải quyết triệt để bài toán tự động hóa cho SME bằng cách kết hợp khoa học dữ liệu (Data Science) "
        "với nghệ thuật Trực quan hóa dữ liệu (Data Visualization). Báo cáo tập trung thực hiện 4 mục tiêu cốt lõi:"
    )
    add_bullet("Nghiên cứu định lượng 180 tác vụ công việc thuộc 8 ngành nghề SME dựa trên bộ dữ liệu WORKBank.", "Mục tiêu 1:")
    add_bullet("Thiết lập công thức toán học chuẩn hóa Min-Max để tính toán Điểm Ưu Tiên (Priority Score) và ROI tiết kiệm chi phí tài chính.", "Mục tiêu 2:")
    add_bullet("Giải mã cấu trúc tâm lý nhân sự thông qua 4 vùng cảm xúc (Sẵn sàng, Lo lắng, Mâu thuẫn, Gắn bó) để xây dựng kịch bản quản trị thay đổi.", "Mục tiêu 3:")
    add_bullet("Xây dựng Dashboard Business Intelligence chuyên nghiệp (phong cách Power BI, Notion, Stripe) tích hợp Trợ Lý AI Advisor giải đáp 24/7.", "Mục tiêu 4:")

    # ═══════════════════════════════════════════════════════════════
    # CHƯƠNG II: BỘ DỮ LIỆU WORKBANK & GIẢI THÍCH CHI TIẾT TỪNG METRICS
    # ═══════════════════════════════════════════════════════════════
    add_h1("CHƯƠNG II: BỘ DỮ LIỆU WORKBANK & GIẢI THÍCH CHI TIẾT TỪNG METRICS")

    add_h2("2.1. Giới Thiệu Bộ Dữ Liệu WORKBank Dataset (arXiv:2506.06576v3)")
    add_p(
        "Nghiên cứu sử dụng dữ liệu thực nghiệm 100% từ công trình khoa học WORKBank do các giáo sư và nhà nghiên cứu từ "
        "Đại học Harvard và Đại học Stanford công bố. Bộ dữ liệu được tích hợp từ 3 nguồn CSV gốc:"
    )

    table_data = [
        ["Tên File Dữ Liệu Gốc", "Số Bản Ghi", "Trường Kế Thừa Cốt Lõi", "Ý Nghĩa Chuyên Môn Trong Phân Tích"],
        ["domain_worker_desires.csv", "180 tác vụ", "task_statement, Desire Score (D), Enjoyment, Security", "Đo lường mức độ mong muốn tự động hóa & tâm lý NLĐ"],
        ["expert_rated_technological_capability.csv", "180 tác vụ", "task_statement, Capability Score (C)", "Đánh giá mức độ phát triển công nghệ AI Agent từ chuyên gia"],
        ["task_statement_with_metadata.csv", "180 tác vụ", "task_statement, occupation, time_hrs, is_sensitive", "Danh mục tác vụ, thời gian thực hiện & từ khóa nhạy cảm"]
    ]

    t = doc.add_table(rows=len(table_data), cols=4)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for r_idx, row in enumerate(table_data):
        for c_idx, val in enumerate(row):
            cell = t.cell(r_idx, c_idx)
            cell.text = val
            set_cell_margins(cell, top=120, bottom=120, left=150, right=150)
            if r_idx == 0:
                set_cell_background(cell, "1E3A8A")
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
            else:
                set_cell_background(cell, "F8FAFC" if r_idx % 2 == 1 else "FFFFFF")

    add_caption("Bảng 1: Cấu trúc chi tiết các thuộc tính bộ dữ liệu khoa học WORKBank")

    add_h2("2.2. Giải Thích Chi Tiết Từng Metric & Chỉ Số Định Lượng Trong Hệ Thống")
    add_p(
        "Để giúp nhà quản trị SME biến các con số khô khan thành định hướng hành động chuẩn xác, "
        "hệ thống xây dựng 7 Metric định lượng cốt lõi với công thức và ý nghĩa kinh doanh cụ thể như sau:"
    )

    add_h3("Metric 1: Normalized Capability Score (Công nghệ AI - C)")
    add_p(
        "• Công thức chuẩn hóa: C = (c - c_min) / (c_max - c_min), với c là điểm đánh giá gốc từ các chuyên gia AI Stanford.\n"
        "• Thang đo: Từ 0.00 (Hoàn toàn không thể tự động hóa) đến 1.00 (AI thực hiện hoàn hảo).\n"
        "• Ý nghĩa: Đo lường mức độ phát triển hiện tại của công nghệ LLM/AI Agent đối với từng tác vụ cụ thể."
    )

    add_h3("Metric 2: Normalized Desire Score (Mong muốn nhân sự - D)")
    add_p(
        "• Công thức chuẩn hóa: D = (d - d_min) / (d_max - d_min), với d là điểm khảo sát mong muốn từ người lao động thực tế.\n"
        "• Thang đo: Từ 0.00 (Nhân viên kiên quyết giữ lại tác vụ) đến 1.00 (Khao khát chuyển giao hoàn toàn cho AI).\n"
        "• Ý nghĩa: Đánh giá mức độ đồng thuận tâm lý của người lao động đối với việc ứng dụng công nghệ."
    )

    add_h3("Metric 3: Priority Score (Điểm Ưu Tiên Chiến Lược - P)")
    add_p(
        "• Công thức toán học: P = (w1 * C + w2 * D) * 100 với trọng số mặc định w1 = 0.5, w2 = 0.5.\n"
        "• Thang điểm: 0 – 100 điểm.\n"
        "• Ý nghĩa: Chỉ số tổng hợp dung hòa cả hai góc độ Năng lực công nghệ (C) và Tâm lý con người (D). Tác vụ có điểm P càng cao thì ưu tiên triển khai càng sớm."
    )

    add_h3("Metric 4: Estimated Annual ROI Savings ($/năm)")
    add_p(
        "• Công thức tài chính: ROI ($/năm) = Task_Time (giờ/tuần) * 52 tuần * Hourly_Salary ($/giờ) * C.\n"
        "• Thang đo: Giá trị tiền tệ ($ USD / năm).\n"
        "• Ý nghĩa: Quy đổi thời gian làm việc tiết kiệm được thành số tiền tài chính thực tế mà SME cắt giảm được hàng năm nếu tự động hóa tác vụ ($1.85M/năm tổng số)."
    )

    add_h3("Metric 5: Strategic Risk Flag (Phân loại rủi ro 4 vùng)")
    add_p(
        "• Phân loại dựa trên ngưỡng 0.5:\n"
        "  - 🟢 Thí điểm Ngay (Ideal Pilot): C ≥ 0.5 & D ≥ 0.5 (95 tác vụ, ROI $938,377/năm).\n"
        "  - 🟡 Phản kháng Nội bộ (Internal Resistance): C ≥ 0.5 & D < 0.5 (16 tác vụ, ROI $252,905/năm).\n"
        "  - 🟠 Kỳ vọng vượt Năng lực (Over Expectation): C < 0.5 & D ≥ 0.5 (0 tác vụ).\n"
        "  - ⚪ Ưu tiên Thấp (Low Priority): C < 0.5 & D < 0.5 (69 tác vụ)."
    )

    add_h3("Metric 6: Sensitivity Flag (Tác vụ nhạy cảm - 34 Tác vụ)")
    add_p(
        "• Định nghĩa: Các tác vụ chứa từ khóa nhạy cảm về tài chính, ngân sách, hợp đồng pháp lý, thông tin nhân sự.\n"
        "• Ý nghĩa quản trị: Bắt buộc thiết lập cơ chế phê duyệt chữ ký số (Human Approval Loop), AI chỉ đóng vai trò Trợ lý (Copilot)."
    )

    add_h3("Metric 7: Worker Psychology Scores (Enjoyment vs Security)")
    add_p(
        "• Enjoyment Score: Mức độ yêu thích và gắn kết cảm xúc của nhân sự đối với tác vụ.\n"
        "• Security Score: Mức độ lo ngại bị sa thải hoặc giảm thu nhập khi AI xuất hiện."
    )

    # ═══════════════════════════════════════════════════════════════
    # CHƯƠNG III: TRỰC QUAN HÓA DỮ LIỆU & PHÂN TÍCH INSIGHT TỪNG BIỂU ĐỒ
    # ═══════════════════════════════════════════════════════════════
    add_h1("CHƯƠNG III: TRỰC QUAN HÓA DỮ LIỆU & PHÂN TÍCH INSIGHT TỪNG BIỂU ĐỒ")

    # ── 3.1. BIỂU ĐỒ 1: GAP ANALYSIS ──
    add_h2("3.1. Biểu Đồ 1: Phân Tích Khoảng Cách Kỳ Vọng (D) vs Năng Lực AI (C)")
    add_image_if_exists("data/reports/charts/chart_1_gap_analysis.png", "Hình 1: Biểu đồ Diverging Bar đại diện cho Khoảng cách Gap = D - C theo từng Vị trí SME")

    add_h3("Cấu Trúc & Cách Đọc Biểu Đồ:")
    add_p(
        "Biểu đồ Diverging Bar Chart thể hiện chênh lệch giữa chỉ số Mong muốn tự động hóa của người lao động (D) "
        "và Năng lực sẵn có của công nghệ AI (C) trên 18 vị trí SME tiêu biểu. "
        "Trục hoành biểu diễn giá trị Gap (%) với trục trung tâm tại gốc 0%. "
        "Thanh bar màu Đỏ (Gap < 0) phản ánh tình trạng AI đã sẵn sàng nhưng nhân sự chưa sẵn lòng đón nhận. "
        "Thanh bar màu Xanh lá (Gap > 0) đại diện cho tình trạng nhân sự kỳ vọng cao nhưng công nghệ AI chưa đáp ứng đủ."
    )

    add_h3("Bằng Chứng Số Liệu Thống Kê & Phân Tích Insight:")
    add_bullet("Nhóm vị trí Kế toán & Thuế, Quản trị dữ liệu có Gap âm lớn nhất (lần lượt đạt -34.5% và -9.0%). Điều này chứng minh công nghệ LLM hiện tại đã tự động hóa cực kỳ tốt các tác vụ tính toán sổ sách, nhưng nhân sự nghiệp vụ vẫn giữ tâm lý e ngại lo sợ bị mất kiểm soát và chịu trách nhiệm pháp lý.", "1. Nhóm Phản Kháng Công Nghệ (Gap Âm):")
    add_bullet("Nhóm Thiết kế đồ họa, Chăm sóc khách hàng có Gap dương vượt trội (+24.5% và +21.3%). Nhân sự các vị trí này chịu áp lực khối lượng công việc lặp lại rất lớn nên rất khao khát có AI Agent hỗ trợ giải phóng thời gian thủ công.", "2. Nhóm Vượt Kỳ Vọng (Gap Dương):")

    add_h3("Đề Xuất Khuyến Nghị Cho Quản Trị SME:")
    add_bullet("Đối với nhóm Gap âm (Kế toán, Quản trị dữ liệu): Tuyệt đối không áp đặt sa thải hoặc ép buộc. Áp dụng mô hình 'Copilot Đồng Hành' — vị trí nhân sự vẫn làm chủ, AI đóng vai trò Trợ lý gợi ý.", "• Chiến lược 1:")
    add_bullet("Đối với nhóm Gap dương (CSKH, Thiết kế): Ưu tiên trang bị các công cụ RAG Agent, sinh nội dung tự động để giải tỏa ngay áp lực công việc hàng ngày.", "• Chiến lược 2:")

    # ── 3.2. BIỂU ĐỒ 2: STRATEGIC QUADRANT ──
    add_h2("3.2. Biểu Đồ 2: Ma Trận 4 Góc Phần Tư Phân Loại Tác Vụ Chiến Lược")
    add_image_if_exists("data/reports/charts/chart_2_strategic_quadrant.png", "Hình 2: Ma trận Scatter 4 Góc phần tư phân loại 180 tác vụ theo Năng lực (C) và Mong muốn (D)")

    add_h3("Cấu Trúc & Cách Đọc Biểu Đồ:")
    add_p(
        "Ma trận Scatter Plot 2 chiều sử dụng Trục X biểu diễn Năng lực AI (C) và Trục Y biểu diễn Mong muốn (D). "
        "Kích thước các bong bóng tỉ lệ thuận với giá trị tiết kiệm ROI ($/năm). "
        "Ma trận chia không gian phân tích thành 4 vùng phân bổ nguồn lực dựa trên ngưỡng phân định 0.5."
    )

    add_h3("Bằng Chứng Số Liệu Thống Kê & Insight Chuyên Sâu:")
    add_bullet("Chiếm 52.8% (95 tác vụ, ROI ước tính đạt $938,377/năm). Các tác vụ này có cả năng lực AI cao và sự đồng thuận của nhân sự.", "🟢 Vùng Thí Điểm Ngay (Ideal Pilot):")
    add_bullet("Chiếm 8.9% (16 tác vụ, ROI ước tính $252,905/năm). AI làm rất tốt nhưng chịu sự chống đối từ nội bộ.", "🟡 Vùng Phản Kháng Nội Bộ (Internal Resistance):")
    add_bullet("Chiếm 0.0% (0 tác vụ). Hiện tại không có tác vụ nào nhân viên muốn tự động hóa mà AI hoàn toàn bó tay.", "🟠 Vùng Kỳ Vọng Vượt Năng Lực (Over Expectation):")
    add_bullet("Chiếm 38.3% (69 tác vụ). Các tác vụ đòi hỏi sự sáng tạo nghệ thuật hoặc có mức độ ưu tiên thấp.", "⚪ Vùng Ưu Tiên Thấp (Low Priority):")

    add_h3("Đề Xuất Khuyến Nghị Phân Bổ Nguồn Lực:")
    add_p(
        "Doanh nghiệp SME cần tập trung 80% ngân sách và nhân lực vào 95 tác vụ thuộc vùng Thí điểm Ngay để nhanh chóng gặt hái thành quả "
        "(Quick Wins) với giá trị $938,377/năm, tạo niềm tin cho toàn bộ tổ chức trước khi mở rộng sang vùng Phản kháng."
    )

    # ── 3.3. BIỂU ĐỒ 3: WORKER PSYCHOLOGY ──
    add_h2("3.3. Biểu Đồ 3: Bản Đồ Tâm Lý Nhân Viên & 4 Vùng Cảm Xúc")
    add_image_if_exists("data/reports/charts/chart_3_worker_psychology.png", "Hình 3: Biểu đồ Bong bóng Phân tích Tâm lý Nhân viên (Enjoyment vs Security)")

    add_h3("Cấu Trúc & Cách Đọc Biểu Đồ:")
    add_p(
        "Biểu đồ phân tích mối quan hệ giữa Mức độ Yêu thích Công việc (Enjoyment - Trục X) "
        "và Mức độ Lo ngại An toàn Công việc (Security - Trục Y). Kích thước các bong bóng tỉ lệ với thời gian tác vụ hàng tuần."
    )

    add_h3("Bằng Chứng Số Liệu & Giải Mã Cấu Trúc Cảm Xúc:")
    add_bullet("Tác vụ hành chính lặp lại mà nhân viên không có gắn kết cảm xúc — Đây là đối tượng tự động hóa thuận lợi nhất.", "• Nhóm 'Sẵn Sàng' (26.1%):")
    add_bullet("Nhân sự không thích tác vụ nhưng rất sợ mất việc do áp lực thu nhập kinh tế gia đình.", "• Nhóm 'Lo Lắng' (13.3%):")
    add_bullet("Nhân sự yêu thích công việc nhưng lo ngại sự xuất hiện của AI sẽ làm thay đổi cấu trúc vai trò.", "• Nhóm 'Mâu Thuẫn' (22.8%):")
    add_bullet("Nhân sự yêu thích công việc và hoàn toàn không lo bị sa thải. Đây là nhóm sẽ phản kháng dữ dội nhất khi bị AI 'tước mất' công việc yêu thích.", "• Nhóm 'Gắn Bó' (16.7%):")

    add_h3("Đề Xuất Kịch Bản Quản Trị Thay Đổi (Change Management):")
    add_p(
        "Nhà quản trị SME không được dùng mệnh lệnh hành chính đối với nhóm Gắn bó. Thay vào đó, định vị AI như một 'Trợ lý tăng cường' "
        "giúp nhân sự làm việc nhanh hơn để họ nâng cao tay nghề chuyên môn."
    )

    # ── 3.4. BIỂU ĐỒ 4 & 5: TOP ROI & PRIORITY DISTRIBUTION ──
    add_h2("3.4. Biểu Đồ 4 & 5: Phân Tích ROI Tác Vụ & Mật Độ Điểm Ưu Tiên")
    add_image_if_exists("data/reports/charts/chart_4_top_roi_tasks.png", "Hình 4: Biểu đồ Bar Top 10 Tác vụ Đem lại Giá trị ROI Tiết kiệm Chi phí Cao nhất ($/năm)", width_inches=5.8)
    add_image_if_exists("data/reports/charts/chart_5_priority_distribution.png", "Hình 5: Biểu đồ Phân bố Mật độ Điểm Ưu Tiên Chiến Lược (Priority Score Distribution)", width_inches=5.8)

    add_h3("Insight Từ Top ROI & Mật Độ Điểm:")
    add_p(
        "Phân tích Top 10 tác vụ ROI cao nhất chỉ ra các công việc về 'Nhập liệu văn phòng', 'Lập báo cáo tài chính', 'Đối soát chứng từ' "
        "đóng góp tới hơn 60% tổng giá trị tiết kiệm chi phí. Điểm ưu tiên trung bình đạt 55.9/100, trong đó phân bố tập trung dày đặc "
        "ở phân khúc từ 65 đến 85 điểm — Khẳng định dư địa tự động hóa vô cùng dồi dào cho SME."
    )

    # ═══════════════════════════════════════════════════════════════
    # CHƯƠNG IV: HỆ THỐNG TRỢ LÝ AI CHATBOT TƯ VẤN CHIẾN LƯỢC
    # ═══════════════════════════════════════════════════════════════
    add_h1("CHƯƠNG IV: HỆ THỐNG TRỢ LÝ AI CHATBOT TƯ VẤN CHIẾN LƯỢC (FLOATING AI ADVISOR)")

    add_h2("4.1. Vai Trò Của Chatbot AI Trợ Lý Trong Nền Tảng Business Intelligence")
    add_p(
        "Một trong những rào cản phổ biến nhất khi áp dụng Dashboard trực quan hóa tại các doanh nghiệp SME là tình trạng: "
        "Nhà quản trị xem các biểu đồ nhưng không biết cách tự đọc hiểu số liệu hoặc chưa đủ kinh nghiệm để suy ra hành động cụ thể. "
        "Để giải quyết triệt để bài toán này, hệ thống tích hợp trực tiếp **Trợ Lý AI Advisor Nhúng Nổi (Floating Strategy AI Widget)** "
        "ngay góc phải bên dưới ứng dụng Streamlit."
    )
    add_p(
        "Chatbot AI đóng vai trò như một **Cố vấn Chiến lược Cấp cao (Senior Strategy Advisor) trực tuyến 24/7**, "
        "sẵn sàng giải đáp mọi thắc mắc của người dùng về dữ liệu và biểu đồ hiển thị trên trang."
    )

    add_h2("4.2. Kiến Trúc RAG Tri Thức WORKBank & Nhận Diện Ngữ Cảnh Trang")
    add_p(
        "Chatbot AI không trả lời chung chung mà được vận hành bởi mô hình LLM tiên tiến (Llama-3.3-70B) kết hợp với "
        "kỹ thuật RAG (Retrieval-Augmented Generation) chứa toàn bộ tri thức bài báo nghiên cứu khoa học WORKBank (Stanford & Harvard):"
    )
    add_bullet("Tự động nhận diện trang người dùng đang mở (Overview, Strategy, Employee Insights, Analysis, Recommendations).", "1. Nhận diện ngữ cảnh trang:")
    add_bullet("Trích xuất chỉ số thống kê tức thì theo vị trí công việc (Job Focus) và ngành nghề người dùng đang chọn trên Sidebar.", "2. Tích hợp dữ liệu thời gian thực:")
    add_bullet("Truy vấn nguyên lý khoa học từ bộ dữ liệu WORKBank để trả lời các câu hỏi về rủi ro, rào cản tâm lý và lộ trình triển khai.", "3. RAG Tri thức Khoa học:")

    add_h2("4.3. Thiết Kế Giao Diện Nổi Bật & Tính Năng Hỗ Trợ Người Dùng")
    add_p("Giao diện Chatbot AI được tối ưu hóa thị giác để kích thích tương tác tự nhiên:")
    add_bullet("Thẻ bong bóng nhịp đập 🟢 💡 Thắc mắc về trang này? Hỏi AI Advisor ngay! hiển thị ngay trên nút bấm kèm hiệu ứng float bounce thu hút sự chú ý.", "• Bong bóng hướng dẫn Callout Badge:")
    add_bullet("Nút bấm Gradient Xanh 💬 Hỏi AI Trợ Lý Về Trang Này kèm hiệu ứng phát sáng Pulse Glow nổi bật hẳn trên nền Dashboard.", "• Nút bấm Popover Nổi:")
    add_bullet("Cung cấp 4 phím tắt hỏi nhanh 1-click (Insight trang này?, Rủi ro nhân sự?, Tác vụ thí điểm trước?, Action Plan 4 bước?) giúp người dùng không cần suy nghĩ câu hỏi.", "• Gợi ý 1-click tiện lợi:")

    add_h2("4.4. Khả Năng Giải Đáp Thắc Mắc Thực Tế Cho Doanh Nghiệp")
    add_p("AI Chatbot hỗ trợ trả lời trực tiếp các thắc mắc kinh điển của lãnh đạo SME:")
    add_bullet("'Biểu đồ Gap này có ý nghĩa gì đối với phòng Kế toán doanh nghiệp tôi?' -> AI trích xuất số liệu Gap -34.5% và giải thích rào cản tâm lý.", "• Thắc mắc 1 (Đọc hiểu biểu đồ):")
    add_bullet("'Tại sao tác vụ nhập liệu lại thuộc vùng Thí điểm Ngay và tiết kiệm được bao nhiêu tiền?' -> AI quy đổi giá trị ROI cụ thể ($/năm).", "• Thắc mắc 2 (Tính toán ROI):")
    add_bullet("'Nhân viên phòng Marketing e ngại bị thay thế thì tôi phải xử lý ra sao?' -> AI tư vấn kịch bản quản trị thay đổi theo nhóm cảm xúc Mâu Thuẫn.", "• Thắc mắc 3 (Quản trị nhân sự):")

    # ═══════════════════════════════════════════════════════════════
    # CHƯƠNG V: ĐỀ XUẤT CHIẾN LƯỢC & LỘ TRÌNH THỰC THI CHO SME
    # ═══════════════════════════════════════════════════════════════
    add_h1("CHƯƠNG V: ĐỀ XUẤT CHIẾN LƯỢC & LỘ TRÌNH THỰC THI CHO DOANH NGHIỆP SME")

    add_h2("5.1. Lộ Trình Triển Khai 4 Giai Đoạn (Action Plan Roadmap)")
    add_bullet("Tập trung đào tạo nhận thức, giải tỏa rào cản tâm lý cho nhân sự và xây dựng quy chế an toàn dữ liệu.", "Giai đoạn 1 (Tuần 1–2): Chuẩn bị & Truyền thông:")
    add_bullet("Triển khai ngay 95 tác vụ Quick Wins cho khối Kế toán và CSKH để thu hồi $938,377/năm.", "Giai đoạn 2 (Tuần 3–6): Thí điểm hẹp (Pilot):")
    add_bullet("Tích hợp AI Agent vào luồng vận hành các phòng ban khác, theo dõi hiệu năng thời gian thực.", "Giai đoạn 3 (Tuần 7–10): Mở rộng & Giám sát:")
    add_bullet("Đóng gói quy trình chuẩn Human-AI Hybrid, chuẩn hóa bộ KPI đánh giá năng suất mới.", "Giai đoạn 4 (Từ tuần 11): Chuẩn hóa Quy trình:")

    add_h2("5.2. Quy Che Quản Trị Rủi Ro & Chữ Ký Phê Duyệt (Human Approval Loop)")
    add_p(
        "Hệ thống phát hiện 34 tác vụ chứa từ khóa nhạy cảm (tài chính, ngân sách, hợp đồng pháp lý, thông tin nhân sự). "
        "Doanh nghiệp SME bắt buộc thiết lập cơ chế phê duyệt chữ ký số (Human-in-the-loop). AI Agent chỉ đóng vai trò "
        "Trợ lý soạn thảo, mọi giao dịch tài chính hoặc quyết định nhân sự đều cần sự phê duyệt của Quản lý."
    )

    add_h2("5.3. Kết Luận Tổng Quan Báo Cáo Cuối Kỳ")
    add_p(
        "Báo cáo bài tập nhóm cuối kỳ đã chứng minh tính đúng đắn và hiệu quả vượt trội của việc áp dụng Trực quan hóa dữ liệu định lượng "
        "kết hợp với Trợ Lý AI Chatbot trong quản trị doanh nghiệp SME. Sự kết hợp giữa bộ dữ liệu khoa học WORKBank, các mô hình toán học chuẩn mực, "
        "giao diện Dashboard BI đẳng cấp và AI Advisor 24/7 chính là giải pháp toàn diện giúp các SME tự tin hội nhập vào kỷ nguyên số."
    )

    doc.save(output_file)
    print(f"Comprehensive final report successfully saved to: {output_file}")

if __name__ == "__main__":
    create_comprehensive_report()
