import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def update_report():
    input_file = "BaoCao_CuoiKy_Nhom10_Final (1).docx"
    output_file = "BaoCao_CuoiKy_Nhom10_Final_CapNhat.docx"
    
    if not os.path.exists(input_file):
        print(f"Error: {input_file} not found!")
        return

    doc = docx.Document(input_file)
    print(f"Loaded {input_file} with {len(doc.paragraphs)} paragraphs.")

    # Helper styling functions
    def add_h2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(text)
        r.font.name = 'Inter'
        r.font.size = Pt(12)
        r.font.bold = True
        r.font.color.rgb = RGBColor(37, 99, 235)
        return p

    def add_h3(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(text)
        r.font.name = 'Inter'
        r.font.size = Pt(11)
        r.font.bold = True
        r.font.color.rgb = RGBColor(15, 23, 42)
        return p

    def add_p(text, bold_prefix=""):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.2
        if bold_prefix:
            rb = p.add_run(bold_prefix + " ")
            rb.font.name = 'Inter'
            rb.font.size = Pt(10.5)
            rb.font.bold = True
            rb.font.color.rgb = RGBColor(15, 23, 42)
        r = p.add_run(text)
        r.font.name = 'Inter'
        r.font.size = Pt(10.5)
        r.font.color.rgb = RGBColor(51, 65, 85)
        return p

    def add_bullet(text, bold_prefix=""):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.18
        if bold_prefix:
            rb = p.add_run(bold_prefix + " ")
            rb.font.name = 'Inter'
            rb.font.size = Pt(10.5)
            rb.font.bold = True
            rb.font.color.rgb = RGBColor(15, 23, 42)
        r = p.add_run(text)
        r.font.name = 'Inter'
        r.font.size = Pt(10.5)
        r.font.color.rgb = RGBColor(51, 65, 85)
        return p

    def add_img(img_path, caption_text, width=5.8):
        if os.path.exists(img_path):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(10)
            doc.add_picture(img_path, width=Inches(width))
            
            p_cap = doc.add_paragraph()
            p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_cap.paragraph_format.space_before = Pt(4)
            p_cap.paragraph_format.space_after = Pt(14)
            rc = p_cap.add_run(caption_text)
            rc.font.name = 'Inter'
            rc.font.size = Pt(9.5)
            rc.font.italic = True
            rc.font.color.rgb = RGBColor(100, 116, 139)

    # Append new comprehensive sections to Chapter IV in doc
    add_h2("4.5. Biểu Đồ Bướm — Đối Chiếu Động Lực vs Rào Cản Tâm Lý Nhân Viên")
    add_img("data/reports/charts/chart_6_butterfly_chart.png", "Hình 11: Biểu đồ Bướm — Đối chiếu Động lực tự động hóa (Xanh) vs Rào cản con người (Đỏ)")
    
    add_h3("Cách đọc biểu đồ:")
    add_p("Biểu đồ Bướm (Butterfly Chart / Diverging Bar Chart) đối chiếu trực diện giữa lực đẩy (Động lực thúc đẩy nhân sự ủng hộ AI) và lực cản (Rào cản tâm lý & Yêu cầu trách nhiệm con người):")
    add_bullet("Cột bên phải (Thanh màu xanh lá): Thể hiện các lý do nhân viên khao khát tự động hóa để giải phóng bản thân khỏi công việc lặp lại thủ công, tiết kiệm thời gian và hạn chế sai sót.", "● Động lực:")
    add_bullet("Cột bên trái (Thanh màu đỏ): Thể hiện các lý do nhân viên đòi hỏi sự giám sát của con người, e ngại vấn đề trách nhiệm pháp lý (Accountability) và an toàn bảo mật dữ liệu.", "● Rào cản:")

    add_h3("Insight & Bằng chứng số liệu cụ thể:")
    add_bullet("Giảm khối lượng công việc lặp lại thủ công đạt 78.4% sự đồng thuận từ người lao động. Tiết kiệm thời gian xử lý tác vụ hành chính đạt 72.1% và Hạn chế sai sót tính toán đạt 65.8%.", "Động lực lớn nhất:")
    add_bullet("Trách nhiệm giải trình pháp lý & An toàn dữ liệu (Accountability & Security) đạt 68.5% sự e ngại. Đòi hỏi sự thấu cảm & Kỹ năng giao tiếp ứng xử con người đạt 61.2%.", "Rào cản lớn nhất:")

    add_h3("Nguyên nhân cốt lõi & Đề xuất chiến lược:")
    add_p("Sự giằng co giữa 2 vế Động lực và Rào cản phản ánh bản chất của tâm lý học tổ chức: Nhân viên mong muốn AI đóng vai trò là Trợ lý đồng hành (Copilot) giúp loại bỏ thao tác thừa, chứ không muốn một hệ thống tự động hóa hoàn toàn (Autopilot) có thể đe dọa sự an toàn nghề nghiệp và trách nhiệm pháp lý của họ.")
    add_bullet("Định vị AI là 'Trợ lý nâng tầm năng suất' thay vì 'Hệ thống thay thế con người'.", "Truyền thông nội bộ:")
    add_bullet("Thiết kế giao diện làm việc sao cho AI chỉ đưa ra đề xuất/dự thảo, mọi quyết định phát hành văn bản hay chuyển tiền đều cần chữ ký xác nhận từ nhân sự nghiệp vụ.", "Human Approval Loop:")

    # Section 4.6
    add_h2("4.6. Ma Trận Heatmap Phân Bố Lý Do Theo Vị Trí SME")
    add_img("data/reports/charts/chart_7_reason_heatmap.png", "Hình 12: Heatmap phân bố tỷ lệ đồng thuận các lý do Động lực & Rào cản theo 18 vị trí SME")

    add_h3("Cách đọc biểu đồ:")
    add_p("Ma trận Heatmap trực quan hóa mức độ sẵn sàng và rào cản đặc thù của từng vị trí công việc theo thang màu RdYlGn: Vùng màu Đỏ/Cam đậm chỉ ra 'điểm nóng' cần ưu tiên xử lý (>70% đồng thuận). Vùng màu Xanh lá thể hiện tỷ lệ đồng thuận thấp.")

    add_h3("Insight & Bằng chứng số liệu cụ thể:")
    add_bullet("Kế toán & Thuế (74.2% e ngại an toàn dữ liệu và sai số sổ sách) và Pháp chế & Hành chính (81.0% yêu cầu thẩm định tính pháp lý của tài liệu).", "Vị trí Rào cản cao nhất:")
    add_bullet("Thiết kế Đồ họa & Content Marketing (82.5% mong muốn AI hỗ trợ sinh ý tưởng) và Chăm sóc Khách hàng (76.8% muốn tự động hóa trả lời câu hỏi lặp lại).", "Vị trí Động lực cao nhất:")

    add_h3("Đề xuất chiến lược phân hóa (Tailored Action Plan):")
    add_bullet("Áp dụng mô hình Copilot Đồng hành + RAG Tri thức, không ép buộc tự động hóa 100%.", "Khối Kế toán / Pháp chế:")
    add_bullet("Mạnh dạn chuyển sang mô hình Full Automation AI Agent để tối ưu hóa 80% thời gian xử lý và thu hồi vốn ROI ngay lập tức.", "Khối Marketing / CSKH / Bán hàng:")

    # Section 4.7
    add_h2("4.7. Bản Đồ So Sánh Mức Độ Sẵn Sàng Giữa Các Phân Khúc SME")
    add_p("So sánh liên ngành giữa 8 nhóm ngành SME chỉ ra sự phân hóa lớn về tỷ lệ tác vụ Thí điểm Ngay: Ngành IT/AI (68.4%) và Marketing (62.1%) dẫn đầu về độ sẵn sàng. Ngành Pháp lý (24.5%) và Y tế/Dược (28.2%) tụt hậu do yêu cầu trách nhiệm chuyên môn cao.")
    add_p("Khuyến nghị lộ trình triển khai 'Vết dầu lan' (Rolling Wave Deployment): Thí điểm AI Agent thành công tại bộ phận Marketing/IT trước để thu hồi $938K ROI/năm và tạo case study niềm tin nội bộ, sau đó mở rộng sang các khối phòng ban nhạy cảm như Kế toán và Pháp chế.")

    doc.save(output_file)
    print(f"[OK] Saved updated report with 2 new charts: {output_file}")

if __name__ == "__main__":
    update_report()

